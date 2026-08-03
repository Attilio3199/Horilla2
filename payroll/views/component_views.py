"""
component_views.py

This module is used to write methods to the component_urls patterns respectively
"""

import json
import math
import operator
from collections import defaultdict
from datetime import date, datetime, timedelta
from itertools import groupby
from urllib.parse import parse_qs

import pandas as pd
from django.apps import apps
from django.conf import settings
from django.contrib import messages
from django.db import connection as _pg_conn
from django.db.models import Sum
from django.http import HttpResponse, HttpResponseRedirect, JsonResponse, QueryDict
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils.html import format_html
from django.utils.translation import gettext_lazy as _
from django.views.decorators.cache import never_cache
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, Side
from openpyxl.utils import get_column_letter

import payroll.models.models
from base.backends import ConfiguredEmailBackend
from base.methods import (
    closest_numbers,
    eval_validate,
    filter_own_records,
    get_key_instances,
    get_next_month_same_date,
    sortby,
)
from base.models import Company
from employee.models import Employee, EmployeeWorkInformation
from horilla.decorators import (
    handle_no_permission,
    hx_request_required,
    login_required,
    owner_can_enter,
    permission_required,
)
from horilla.group_by import group_by_queryset
from horilla.settings.base import HORILLA_DATE_FORMATS
from horilla.http.response import HorillaRedirect
from horilla.methods import dynamic_attr, get_horilla_model_class, get_urlencode

# from leave.models import AvailableLeave
from notifications.signals import notify
from payroll.filters import (
    AllowanceFilter,
    DeductionFilter,
    LoanAccountFilter,
    PayslipFilter,
    PayslipReGroup,
    ReimbursementFilter,
)
from payroll.forms import component_forms as forms
from payroll.methods.deductions import create_deductions, update_compensation_deduction
from payroll.methods.methods import (
    calculate_employer_contribution,
    compute_net_pay,
    compute_salary_on_period,
    paginator_qry,
    save_payslip,
)
from payroll.methods.payslip_calc import (
    calculate_allowance,
    calculate_gross_pay,
    calculate_net_pay_deduction,
    calculate_post_tax_deduction,
    calculate_pre_tax_deduction,
    calculate_tax_deduction,
    calculate_taxable_gross_pay,
)
from payroll.methods.tax_calc import calculate_taxable_amount
from payroll.models.models import (
    Allowance,
    Contract,
    Deduction,
    LoanAccount,
    PayslipControlloRegola,
    PayslipControlloRegolaDestinazione,
    PayslipCorpo,
    Payslip,
    PayslipDizionario,
    PayslipImporti,
    PayslipPresenze,
    Reimbursement,
    ReimbursementMultipleAttachment,
)
from payroll.threadings.mail import MailSendThread

# from asset.models import Asset


def return_none(a, b):
    return None


operator_mapping = {
    "equal": operator.eq,
    "notequal": operator.ne,
    "lt": operator.lt,
    "gt": operator.gt,
    "le": operator.le,
    "ge": operator.ge,
    "icontains": operator.contains,
    "range": return_none,
}


def payroll_calculation(employee, start_date, end_date):
    """
    Calculate payroll components for the specified employee within the given date range.


    Args:
        employee (Employee): The employee for whom the payroll is calculated.
        start_date (date): The start date of the payroll period.
        end_date (date): The end date of the payroll period.


    Returns:
        dict: A dictionary containing the calculated payroll components:
    """

    basic_pay_details = compute_salary_on_period(employee, start_date, end_date)
    if not basic_pay_details:
        return None
    contract = basic_pay_details["contract"]
    contract_wage = basic_pay_details["contract_wage"]
    basic_pay = basic_pay_details["basic_pay"]
    loss_of_pay = basic_pay_details["loss_of_pay"]
    custom_leave_deduction = basic_pay_details.get("custom_leave_deduction", 0.0)
    custom_leave_breakdown = basic_pay_details.get("custom_leave_breakdown", [])
    paid_days = basic_pay_details["paid_days"]
    unpaid_days = basic_pay_details["unpaid_days"]
    partial_pay_days = basic_pay_details.get("partial_pay_days", 0)

    working_days_details = basic_pay_details["month_data"]

    updated_basic_pay_data = update_compensation_deduction(
        employee, basic_pay, "basic_pay", start_date, end_date
    )
    basic_pay = updated_basic_pay_data["compensation_amount"]
    basic_pay_deductions = updated_basic_pay_data["deductions"]

    loss_of_pay_amount = 0
    if not contract.deduct_leave_from_basic_pay:
        loss_of_pay_amount = loss_of_pay
    else:
        basic_pay = basic_pay - loss_of_pay_amount

    kwargs = {
        "employee": employee,
        "start_date": start_date,
        "end_date": end_date,
        "basic_pay": basic_pay,
        "day_dict": working_days_details,
    }
    # basic pay will be basic_pay = basic_pay - update_compensation_amount
    allowances = calculate_allowance(**kwargs)

    # finding the total allowance
    total_allowance = sum(allowance["amount"] for allowance in allowances["allowances"])

    kwargs["allowances"] = allowances
    kwargs["total_allowance"] = total_allowance
    updated_gross_pay_data = calculate_gross_pay(**kwargs)
    gross_pay = updated_gross_pay_data["gross_pay"]
    gross_pay_deductions = updated_gross_pay_data["deductions"]

    kwargs["gross_pay"] = gross_pay
    pretax_deductions = calculate_pre_tax_deduction(**kwargs)
    post_tax_deductions = calculate_post_tax_deduction(**kwargs)

    installments = (
        pretax_deductions["installments"] | post_tax_deductions["installments"]
    )

    taxable_gross_pay = calculate_taxable_gross_pay(**kwargs)
    tax_deductions = calculate_tax_deduction(**kwargs)
    federal_tax = calculate_taxable_amount(**kwargs)

    total_allowance = sum(item["amount"] for item in allowances["allowances"])
    total_pretax_deduction = sum(
        item["amount"] for item in pretax_deductions["pretax_deductions"]
    )
    total_post_tax_deduction = sum(
        item["amount"] for item in post_tax_deductions["post_tax_deductions"]
    )
    total_tax_deductions = sum(
        item["amount"] for item in tax_deductions["tax_deductions"]
    )

    total_deductions = (
        total_pretax_deduction
        + total_post_tax_deduction
        + total_tax_deductions
        + federal_tax
        + loss_of_pay  # 1022
    )

    net_pay = gross_pay - total_deductions
    # loss_of_pay        -> actual lop amount
    # loss_of_pay_amount -> actual lop if deduct from basic-
    #                       pay from contract is enabled
    net_pay = compute_net_pay(
        net_pay=net_pay,
        gross_pay=gross_pay,
        total_pretax_deduction=total_pretax_deduction,
        total_post_tax_deduction=total_post_tax_deduction,
        total_tax_deductions=total_tax_deductions,
        federal_tax=federal_tax,
        loss_of_pay_amount=loss_of_pay_amount,
        loss_of_pay=loss_of_pay,
    )
    updated_net_pay_data = update_compensation_deduction(
        employee, net_pay, "net_pay", start_date, end_date
    )
    net_pay = updated_net_pay_data["compensation_amount"]
    update_net_pay_deductions = updated_net_pay_data["deductions"]

    net_pay_deductions = calculate_net_pay_deduction(
        net_pay,
        post_tax_deductions["net_pay_deduction"],
        **kwargs,
    )
    net_pay_deduction_list = net_pay_deductions["net_pay_deductions"]
    for deduction in update_net_pay_deductions:
        net_pay_deduction_list.append(deduction)
    net_pay = net_pay - net_pay_deductions["net_deduction"]
    payslip_data = {
        "employee": employee,
        "contract_wage": contract_wage,
        "basic_pay": basic_pay,
        "gross_pay": gross_pay,
        "taxable_gross_pay": taxable_gross_pay["taxable_gross_pay"],
        "net_pay": net_pay,
        "allowances": allowances["allowances"],
        "paid_days": paid_days,
        "unpaid_days": unpaid_days,
        "partial_pay_days": partial_pay_days,
        "basic_pay_deductions": basic_pay_deductions,
        "gross_pay_deductions": gross_pay_deductions,
        "pretax_deductions": pretax_deductions["pretax_deductions"],
        "post_tax_deductions": post_tax_deductions["post_tax_deductions"],
        "tax_deductions": tax_deductions["tax_deductions"],
        "net_deductions": net_pay_deduction_list,
        "total_deductions": total_deductions,
        "loss_of_pay": loss_of_pay,
        "custom_leave_deduction": custom_leave_deduction,
        "custom_leave_breakdown": custom_leave_breakdown,
        "federal_tax": federal_tax,
        "start_date": start_date,
        "end_date": end_date,
        "range": f"{start_date.strftime('%b %d %Y')} - {end_date.strftime('%b %d %Y')}",
    }
    data_to_json = payslip_data.copy()
    data_to_json["employee"] = employee.id
    data_to_json["start_date"] = start_date.strftime("%Y-%m-%d")
    data_to_json["end_date"] = end_date.strftime("%Y-%m-%d")
    json_data = json.dumps(data_to_json)

    payslip_data["json_data"] = json_data
    payslip_data["installments"] = installments
    return payslip_data


@login_required
@hx_request_required
def allowances_deductions_tab(request, emp_id):
    """
    Retrieve and render the allowances and deductions applicable to an employee.

    This view function retrieves the active contract, basic pay, allowances, and
    deductions for a specified employee. It filters allowances and deductions
    based on various conditions, including specific employee assignments and
    condition-based rules. The results are then rendered in the allowance and
    deduction tab template.
    """
    user = request.user
    employee_deductions = []
    employee_allowances = []
    employee = Employee.objects.get(id=emp_id)
    if getattr(user, "employee_get", None) != employee and not (
        user.has_perm("payroll.view_allowance")
        and user.has_perm("payroll.view_deduction")
    ):
        return handle_no_permission(request)

    active_contracts = employee.contract_set.filter(contract_status="active").first()
    basic_pay = active_contracts.wage if active_contracts else None
    if basic_pay:
        allowances = (
            Allowance.objects.filter(specific_employees=employee)
            | Allowance.objects.filter(is_condition_based=True).exclude(
                exclude_employees=employee
            )
            | Allowance.objects.filter(include_active_employees=True).exclude(
                exclude_employees=employee
            )
        )

        for allowance in allowances:
            applicable = True
            if allowance.is_condition_based:
                conditions = list(
                    allowance.other_conditions.values_list(
                        "field", "condition", "value"
                    )
                )
                conditions.append(
                    (
                        allowance.field,
                        allowance.condition,
                        allowance.value.lower().replace(" ", "_"),
                    )
                )
                for field, operator, value in conditions:
                    val = dynamic_attr(employee, field)
                    if val is None or not operator_mapping.get(operator)(
                        val, type(val)(value)
                    ):
                        applicable = False
                        break
            if applicable and allowance not in employee_allowances:
                employee_allowances.append(allowance)

        employee_allowances = [
            allowance
            for allowance in employee_allowances
            if operator_mapping.get(allowance.if_condition)(
                basic_pay if allowance.if_choice == "basic_pay" else 0,
                allowance.if_amount,
            )
        ]

        # Find the applicable deductions for the employee
        deductions = (
            Deduction.objects.filter(
                specific_employees=employee,
            )
            | Deduction.objects.filter(
                is_condition_based=True,
            ).exclude(exclude_employees=employee)
            | Deduction.objects.filter(
                include_active_employees=True,
            ).exclude(exclude_employees=employee)
        )
        for deduction in deductions:
            applicable = True
            if deduction.is_condition_based:
                conditions = list(
                    deduction.other_conditions.values_list(
                        "field", "condition", "value"
                    )
                )
                conditions.append(
                    (
                        deduction.field,
                        deduction.condition,
                        deduction.value.lower().replace(" ", "_"),
                    )
                )
                for field, operator, value in conditions:
                    val = dynamic_attr(employee, field)
                    if val is None or not operator_mapping.get(operator)(
                        val, type(val)(value)
                    ):
                        applicable = False
                        break
            if applicable:
                employee_deductions.append(deduction)

    allowance_ids = (
        json.dumps([instance.id for instance in employee_deductions])
        if employee_deductions
        else None
    )
    deduction_ids = (
        json.dumps([instance.id for instance in employee_deductions])
        if employee_deductions
        else None
    )
    context = {
        "active_contracts": active_contracts,
        "basic_pay": basic_pay,
        "allowances": employee_allowances if employee_allowances else None,
        "allowance_ids": allowance_ids,
        "deductions": employee_deductions if employee_deductions else None,
        "deduction_ids": deduction_ids,
        "employee": employee,
    }
    return render(request, "tabs/allowance_deduction-tab.html", context=context)


@login_required
@permission_required("payroll.add_allowance")
def create_allowance(request):
    """
    This method is used to create allowance condition template
    """
    form = forms.AllowanceForm()
    is_htmx = request.headers.get("HX-Request") is not None
    if request.method == "POST":
        form = forms.AllowanceForm(request.POST)
        if form.is_valid():
            form.save()
            form = forms.AllowanceForm()
            messages.success(request, _("Allowance created."))
            if is_htmx:
                response = HttpResponse("", status=200)
                response["HX-Trigger"] = json.dumps(
                    {"reloadPayrollAllowances": {"target": "body"}}
                )
                return response
            return redirect(reverse("view-allowance"))
    template_name = (
        "payroll/common/form_fragment.html" if is_htmx else "payroll/common/form.html"
    )
    return render(
        request,
        template_name,
        {
            "form": form,
            "post_url": request.get_full_path(),
            "back_url": reverse("allowances-list-view"),
        },
    )


@login_required
@permission_required("payroll.view_allowance")
def view_allowance(request):
    """
    This method is used render template to view all the allowance instances
    """

    allowances = payroll.models.models.Allowance.objects.exclude(
        only_show_under_employee=True
    )
    allowance_filter = AllowanceFilter(request.GET)
    allowances = paginator_qry(allowances, request.GET.get("page"))
    allowance_ids = json.dumps([instance.id for instance in allowances.object_list])
    return render(
        request,
        "payroll/allowance/view_allowance.html",
        {
            "allowances": allowances,
            "f": allowance_filter,
            "allowance_ids": allowance_ids,
        },
    )


@login_required
@hx_request_required
def view_single_allowance(request, allowance_id):
    """
    This method is used render template to view the selected allowance instances
    """
    previous_data = get_urlencode(request)
    allowance = Allowance.find(allowance_id)
    allowance_ids_json = request.GET.get("instances_ids")
    context = {
        "allowance": allowance,
    }
    if allowance_ids_json:
        allowance_ids = json.loads(allowance_ids_json)
        previous_id, next_id = closest_numbers(allowance_ids, allowance_id)
        context["next"] = next_id
        context["previous"] = previous_id
        context["allowance_ids"] = allowance_ids
    context["pd"] = previous_data
    return render(
        request,
        "payroll/allowance/view_single_allowance.html",
        context,
    )


@login_required
@hx_request_required
@permission_required("payroll.view_allowance")
def filter_allowance(request):
    """
    Filter and retrieve a list of allowances based on the provided query parameters.
    """
    query_string = request.GET.urlencode()
    allowances = AllowanceFilter(request.GET).qs.exclude(only_show_under_employee=True)
    list_view = "payroll/allowance/list_allowance.html"
    card_view = "payroll/allowance/card_allowance.html"
    template = card_view
    if request.GET.get("view") == "list":
        template = list_view
    allowances = sortby(request, allowances, "sortby")
    allowances = paginator_qry(allowances, request.GET.get("page"))
    allowance_ids = json.dumps([instance.id for instance in allowances.object_list])
    data_dict = parse_qs(query_string)
    get_key_instances(Allowance, data_dict)
    return render(
        request,
        template,
        {
            "allowances": allowances,
            "pd": query_string,
            "filter_dict": data_dict,
            "allowance_ids": allowance_ids,
        },
    )


@login_required
@permission_required("payroll.change_allowance")
def update_allowance(request, allowance_id, **kwargs):
    """
    This method is used to update the allowance
    Args:
        id : allowance instance id
    """
    instance = Allowance.find(allowance_id)
    is_htmx = request.headers.get("HX-Request") is not None
    if not instance:
        return HorillaRedirect(request, message=_("Allowance not found."))
    form = forms.AllowanceForm(instance=instance)
    if request.method == "POST":
        form = forms.AllowanceForm(request.POST, instance=instance)
        if form.is_valid():
            form.save()
            messages.success(request, _("Allowance updated."))
            if is_htmx:
                response = HttpResponse("", status=200)
                response["HX-Trigger"] = json.dumps(
                    {"reloadPayrollAllowances": {"target": "body"}}
                )
                return response
            return redirect(reverse("view-allowance"))
    template_name = (
        "payroll/common/form_fragment.html" if is_htmx else "payroll/common/form.html"
    )
    return render(
        request,
        template_name,
        {
            "form": form,
            "post_url": request.get_full_path(),
            "back_url": reverse("allowances-list-view"),
        },
    )


# @login_required
# @hx_request_required
# @permission_required("payroll.delete_allowance")
# def delete_allowance(request, allowance_id):
#     """
#     This method is used to delete the allowance instance
#     """
#     target = request.META.get("HTTP_HX_TARGET")


#     try:
#         allowance = payroll.models.models.Allowance.objects.filter(
#             id=allowance_id
#         ).first()
#         if allowance:
#             # allowance.delete()
#             messages.success(request, _("Allowance deleted successfully"))
#         else:
#             messages.error(request, _("Allowance not found"))

#     except ValidationError as validation_error:
#         messages.error(
#             request, _("Validation error occurred while deleting the allowance")
#         )
#         messages.error(request, str(validation_error))
#     except Exception as exception:
#         messages.error(request, _("An error occurred while deleting the allowance"))
#         messages.error(request, str(exception))
#     if target and target == "allowance_id":
#         return redirect(reverse("allowances-list-view"))
#         # return HttpResponse("<script>location.reload();</script>")
#     if target and target == "allowance_tab_id":
#         # return redirect(reverse("allowance-tab-list"))
#         return HttpResponse("<script>location.reload();</script>")

#     if (
#         request.path.split("/")[2] == "delete-employee-allowance"
#         or not payroll.models.models.Allowance.objects.filter()
#     ):
#         return return HorillaRedirect(request)
#     return redirect(filter_allowance)


@login_required
@hx_request_required
@permission_required("payroll.delete_allowance")
def delete_allowance(request, allowance_id, emp_id=None):
    target = request.META.get("HTTP_HX_TARGET")
    instances_ids = request.GET.get("instances_ids")
    next_instance = None
    instances_list = None
    if instances_ids:
        instances_list = json.loads(instances_ids)
        previous_instance, next_instance = closest_numbers(instances_list, allowance_id)
        instances_list.remove(allowance_id)
    allowance = payroll.models.models.Allowance.objects.filter(id=allowance_id).first()
    if allowance:
        allowance.delete()
        messages.success(request, _("Allowance deleted successfully"))
    else:
        messages.error(request, _("Allowance not found"))

    paths = {
        "payroll-deduction-container": f"/payroll/filter_allowance?{request.GET.urlencode()}",
        "allowance_tab_id": f"/payroll/allowance-tab-list/{emp_id}?deleted=true",
        "allowance_id": "/payroll/allowances-list-view/",
        "allowance_card": "/payroll/allowances-card-view/",
        "genericModalBody": f"/payroll/allowance-detail-view/{next_instance}?instance_ids={instances_list}&deleted=true",
    }
    http_hx_target = request.META.get("HTTP_HX_TARGET")
    redirected_path = paths.get(http_hx_target)
    if http_hx_target:
        if (
            http_hx_target == "payroll-deduction-container"
            and not Deduction.objects.filter()
        ):
            return HorillaRedirect(request)
        if redirected_path:
            return redirect(redirected_path)

    default_redirect = (
        request.path if http_hx_target else request.META.get("HTTP_REFERER", "/")
    )
    return HttpResponseRedirect(default_redirect)


@login_required
@permission_required("payroll.add_deduction")
def create_deduction(request):
    """
    This method is used to create deduction
    """
    form = forms.DeductionForm()
    is_htmx = request.headers.get("HX-Request") is not None
    if request.method == "POST":
        form = forms.DeductionForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, _("Deduction created."))
            if is_htmx:
                response = HttpResponse("", status=200)
                response["HX-Trigger"] = json.dumps(
                    {"reloadPayrollDeductions": {"target": "body"}}
                )
                return response
            return redirect(reverse("view-deduction"))
    template_name = (
        "payroll/common/form_fragment.html" if is_htmx else "payroll/common/form.html"
    )
    return render(
        request,
        template_name,
        {
            "form": form,
            "post_url": request.get_full_path(),
            "back_url": reverse("deduction-view-list"),
        },
    )


@login_required
@permission_required("payroll.view_allowance")
def view_deduction(request):
    """
    This method is used render template to view all the deduction instances
    """

    deductions = Deduction.objects.exclude(only_show_under_employee=True)
    deduction_filter = DeductionFilter(request.GET)
    deductions = paginator_qry(deductions, request.GET.get("page"))
    deduction_ids = json.dumps([instance.id for instance in deductions.object_list])
    return render(
        request,
        "payroll/deduction/view_deduction.html",
        {
            "deductions": deductions,
            "f": deduction_filter,
            "deduction_ids": deduction_ids,
        },
    )


@login_required
@hx_request_required
def view_single_deduction(request, deduction_id):
    """
    Render template to view a single deduction instance with navigation.
    """
    previous_data = get_urlencode(request)
    deduction = Deduction.objects.filter(id=deduction_id).first()
    context = {"deduction": deduction, "pd": previous_data}

    # Handle deduction IDs and navigation
    deduction_ids_json = request.GET.get("instances_ids")
    if deduction_ids_json:
        deduction_ids = json.loads(deduction_ids_json)
        context["previous"], context["next"] = closest_numbers(
            deduction_ids, deduction_id
        )
        context["deduction_ids"] = deduction_ids

    # Determine htmx load URL and target
    HTTP_REFERER = request.META.get("HTTP_REFERER", "")
    referer_parts = HTTP_REFERER.rstrip("/").split("/")

    if "view-deduction" in referer_parts:
        context.update(
            {
                "load_hx_url": f"/payroll/filter-deduction?{previous_data}",
                "load_hx_target": "#payroll-deduction-container",
            }
        )
    elif referer_parts[-2:] == ["employee-view", str(referer_parts[-1])]:
        try:
            context.update(
                {
                    "load_hx_url": f"/payroll/allowances-deductions-tab/{int(referer_parts[-1])}",
                    "load_hx_target": "#allowance_deduction",
                }
            )
        except ValueError:
            pass
    elif HTTP_REFERER.endswith("employee-profile/"):
        context.update(
            {
                "load_hx_url": f"/payroll/allowances-deductions-tab/{request.user.employee_get.id}",
                "load_hx_target": "#allowance_deduction",
            }
        )
    else:
        context.update({"load_hx_url": None, "load_hx_target": None})

    return render(request, "payroll/deduction/view_single_deduction.html", context)


@login_required
@hx_request_required
@permission_required("payroll.view_allowance")
def filter_deduction(request):
    """
    This method is used search the deduction
    """
    query_string = request.GET.urlencode()
    deductions = DeductionFilter(request.GET).qs.exclude(only_show_under_employee=True)
    list_view = "payroll/deduction/list_deduction.html"
    card_view = "payroll/deduction/card_deduction.html"
    template = card_view
    if request.GET.get("view") == "list":
        template = list_view
    deductions = sortby(request, deductions, "sortby")
    deductions = paginator_qry(deductions, request.GET.get("page"))
    deduction_ids = json.dumps([instance.id for instance in deductions.object_list])
    data_dict = parse_qs(query_string)
    get_key_instances(Deduction, data_dict)
    return render(
        request,
        template,
        {
            "deductions": deductions,
            "pd": query_string,
            "filter_dict": data_dict,
            "deduction_ids": deduction_ids,
        },
    )


@login_required
@permission_required("payroll.change_deduction")
def update_deduction(request, deduction_id, **kwargs):
    """
    This method is used to update the deduction instance
    """
    instance = Deduction.find(deduction_id)
    is_htmx = request.headers.get("HX-Request") is not None
    if not instance:
        return HorillaRedirect(request, message=_("Deduction not found."))
    form = forms.DeductionForm(instance=instance)
    if request.method == "POST":
        form = forms.DeductionForm(request.POST, instance=instance)
        if form.is_valid():
            form.save()
            messages.success(request, _("Deduction updated."))
            if is_htmx:
                response = HttpResponse("", status=200)
                response["HX-Trigger"] = json.dumps(
                    {"reloadPayrollDeductions": {"target": "body"}}
                )
                return response
            return redirect(reverse("view-deduction"))
    template_name = (
        "payroll/common/form_fragment.html" if is_htmx else "payroll/common/form.html"
    )
    return render(
        request,
        template_name,
        {
            "form": form,
            "post_url": request.get_full_path(),
            "back_url": reverse("deduction-view-list"),
        },
    )


@login_required
@hx_request_required
@permission_required("payroll.delete_deduction")
def delete_deduction(request, deduction_id, emp_id=None):
    instances_ids = request.GET.get("instances_ids")
    next_instance = None
    instances_list = None
    previous_data = ""
    if instances_ids:
        previous_data = get_urlencode(request)
        instances_list = json.loads(instances_ids)
        previous_instance, next_instance = closest_numbers(instances_list, deduction_id)
        instances_list.remove(deduction_id)
    deduction = Deduction.objects.filter(id=deduction_id).first()
    if deduction:
        deduction.delete()
        messages.success(request, _("Deduction deleted successfully"))
    else:
        messages.error(request, _("Deduction not found"))

    paths = {
        "deduct-container": f"/payroll/deduction-view-list?{request.GET.urlencode()}",
        "payroll-deduction-container": f"/payroll/filter-deduction?{request.GET.urlencode()}",
        "allowance_deduction": f"/employee/allowances-deductions-tab/{emp_id}",
        "deduct-div": f"/payroll/deduction-tab-list/{emp_id}?deleted=true",
        "objectDetailsModalTarget": f"/payroll/single-deduction-view/{next_instance}?instances_ids={instances_list}",
        "genericModalBody": f"/payroll/deduction-detail-view/{next_instance}?instance_ids={instances_list}&deleted=true",
    }
    http_hx_target = request.META.get("HTTP_HX_TARGET")
    redirected_path = paths.get(http_hx_target)
    if http_hx_target:
        if (
            http_hx_target == "payroll-deduction-container"
            and not Deduction.objects.filter()
        ):
            return HorillaRedirect(request)
        if redirected_path:
            return redirect(redirected_path)

    default_redirect = (
        request.path if http_hx_target else request.META.get("HTTP_REFERER", "/")
    )
    return HttpResponseRedirect(default_redirect)


def get_month_start_end(year):
    start_end_dates = []
    for month in range(1, 13):
        # Start date is the first day of the month
        start_date = date(year, month, 1)

        # Calculate the last day of the month
        if month == 12:  # December
            end_date = date(year, 12, 31)
        else:
            next_month = date(year, month + 1, 1)
            end_date = next_month - timedelta(days=1)

        start_end_dates.append((start_date, end_date))
    return start_end_dates


@login_required
@permission_required("payroll.add_payslip")
def generate_payslip(request):
    """
    Generate payslips for selected employees within a specified date range.

    Requires the user to be logged in and have the 'payroll.add_payslip' permission.

    """
    if (
        request.META.get("HTTP_HX_REQUEST")
        and request.META.get("HTTP_HX_TARGET") == "objectCreateModalTarget"
    ):
        bulk_form = forms.GeneratePayslipForm()
        return render(
            request,
            "payroll/payslip/bulk_create_payslip.html",
            {"bulk_form": bulk_form},
        )
    payslips = []
    json_data = []
    form = forms.GeneratePayslipForm()
    if request.method == "POST":
        form = forms.GeneratePayslipForm(request.POST)
        if form.is_valid():
            instances = []
            employees = form.cleaned_data["employee_id"]
            start_date = form.cleaned_data["start_date"]
            end_date = form.cleaned_data["end_date"]

            group_name = form.cleaned_data["group_name"]
            emp_count = employees.count()
            for employee in employees:
                contract = Contract.objects.filter(
                    employee_id=employee, contract_status="active"
                ).first()
                if start_date < contract.contract_start_date:
                    start_date = contract.contract_start_date

                if end_date < start_date:
                    messages.error(
                        request, _(f"{employee}'s contract has not started yet.")
                    )
                    emp_count -= 1
                    continue

                payslip = payroll_calculation(employee, start_date, end_date)
                payslips.append(payslip)
                json_data.append(payslip["json_data"])

                payslip["payslip"] = payslip
                data = {}
                data["employee"] = employee
                data["group_name"] = group_name
                data["start_date"] = payslip["start_date"]
                data["end_date"] = payslip["end_date"]
                data["status"] = "draft"
                data["contract_wage"] = payslip["contract_wage"]
                data["basic_pay"] = payslip["basic_pay"]
                data["gross_pay"] = payslip["gross_pay"]
                data["deduction"] = payslip["total_deductions"]
                data["net_pay"] = payslip["net_pay"]
                data["pay_data"] = json.loads(payslip["json_data"])
                calculate_employer_contribution(data)
                data["installments"] = payslip["installments"]
                instance = save_payslip(**data)
                instances.append(instance)
                notify.send(
                    request.user.employee_get,
                    recipient=employee.employee_user_id,
                    verb="Payslip has been generated for you.",
                    verb_ar="تم إصدار كشف راتب لك.",
                    verb_de="Gehaltsabrechnung wurde für Sie erstellt.",
                    verb_es="Se ha generado la nómina para usted.",
                    verb_fr="La fiche de paie a été générée pour vous.",
                    redirect=reverse(
                        "view-created-payslip", kwargs={"payslip_id": instance.id}
                    ),
                    icon="close",
                )
            messages.success(request, f"{emp_count} payslip saved as draft")
            return redirect(
                f"/payroll/view-payslip/?group_by=group_name&active_group={group_name}"
            )

    return render(request, "payroll/common/form.html", {"form": form})


@login_required
@hx_request_required
def check_contract_start_date(request):
    """
    Check if the employee's contract start date is after the provided payslip start date.
    """

    employee_id = request.GET.get("employee_id")
    start_date = request.GET.get("start_date")

    contract = Contract.objects.filter(
        employee_id=employee_id, contract_status="active"
    ).first()

    if not contract or start_date >= str(contract.contract_start_date):
        return HttpResponse("")

    title_message = _(
        "When this payslip is run, the payslip start date will be updated to match the employee contract start date."
    )
    text_content = _("Employee Contract Start Date")

    return HttpResponse(
        format_html(
            """
        <div id='messageDiv' style='background-color: hsl(48, 100%, 94%);
            border: 1px solid hsl(46, 97%, 88%);
            border-radius: 18px; padding:5px; font-weight: bold; display: flex;'>
            {text_content}: {contract_start_date}
            <img style='width: 20px; height: 20px; cursor: pointer;'
                src='/static/images/ui/info.png' class='ml-2' title='{title_message}'>
        </div>
        """,
            text_content=text_content,
            contract_start_date=contract.contract_start_date,
            title_message=title_message,
        )
    )


@login_required
@hx_request_required
@permission_required("payroll.add_payslip")
def create_payslip(request, new_post_data=None):
    """
    Create a payslip for an employee.

    This method is used to create a payslip for an employee based on the provided form data.

    Args:
        request: The HTTP request object.

    Returns:
        A rendered HTML template for the payslip creation form.
    """
    if new_post_data:
        request.POST = new_post_data

    form = forms.PayslipForm()

    if request.method == "POST":
        employee_id = request.POST.get("employee_id")
        start_date = (
            datetime.strptime(request.POST.get("start_date"), "%Y-%m-%d").date()
            if isinstance(request.POST.get("start_date"), str)
            else request.POST.get("start_date")
        )

        if employee_id and start_date:
            contract = Contract.objects.filter(
                employee_id=employee_id, contract_status="active"
            ).first()

            if contract and start_date < contract.contract_start_date:
                new_post_data = request.POST.copy()
                new_post_data["start_date"] = contract.contract_start_date
                request.POST = new_post_data
        form = forms.PayslipForm(request.POST)
        if form.is_valid():
            employee = form.cleaned_data["employee_id"]
            start_date = form.cleaned_data["start_date"]
            end_date = form.cleaned_data["end_date"]
            payslip = Payslip.objects.filter(
                employee_id=employee, start_date=start_date, end_date=end_date
            ).first()

            if form.is_valid():
                employee = form.cleaned_data["employee_id"]
                start_date = form.cleaned_data["start_date"]
                end_date = form.cleaned_data["end_date"]
                payslip_data = payroll_calculation(employee, start_date, end_date)
                payslip_data["payslip"] = payslip
                data = {}
                data["employee"] = employee
                data["start_date"] = payslip_data["start_date"]
                data["end_date"] = payslip_data["end_date"]
                data["status"] = (
                    "draft"
                    if request.GET.get("status") is None
                    else request.GET["status"]
                )
                data["contract_wage"] = payslip_data["contract_wage"]
                data["basic_pay"] = payslip_data["basic_pay"]
                data["gross_pay"] = payslip_data["gross_pay"]
                data["deduction"] = payslip_data["total_deductions"]
                data["net_pay"] = payslip_data["net_pay"]
                data["pay_data"] = json.loads(payslip_data["json_data"])
                calculate_employer_contribution(data)
                data["installments"] = payslip_data["installments"]
                payslip_data["instance"] = save_payslip(**data)
                form = forms.PayslipForm()
                messages.success(request, _("Payslip Saved"))
                payslip = payslip_data["instance"]
                notify.send(
                    request.user.employee_get,
                    recipient=employee.employee_user_id,
                    verb="Payslip has been generated for you.",
                    verb_ar="تم إصدار كشف راتب لك.",
                    verb_de="Gehaltsabrechnung wurde für Sie erstellt.",
                    verb_es="Se ha generado la nómina para usted.",
                    verb_fr="La fiche de paie a été générée pour vous.",
                    redirect=reverse(
                        "view-created-payslip", kwargs={"payslip_id": payslip.pk}
                    ),
                    icon="close",
                )
                return HorillaRedirect(
                    request,
                    redirect_to=reverse(
                        "view-payslip", kwargs={"payslip_id": payslip.pk}
                    ),
                )
    return render(
        request,
        "payroll/payslip/create_payslip.html",
        {"individual_form": form},
    )


@login_required
@hx_request_required
@permission_required("payroll.add_payslip")
def validate_start_date(request):
    """
    This method to validate the contract start date and the pay period start date
    """
    end_datetime = None
    start_datetime = None
    valid = True
    errors = []
    start_date = request.GET.get("start_date")
    end_date = request.GET.get("end_date")
    try:
        employee_id = [
            int(e) for e in request.GET.getlist("employee_id") if e.isdigit()
        ]
    except:
        return HorillaRedirect(request, message=_("Invalid Request"))

    if start_date:
        start_datetime = datetime.strptime(start_date, "%Y-%m-%d").date()
    if end_date:
        end_datetime = datetime.strptime(end_date, "%Y-%m-%d").date()
    for emp_id in employee_id:
        contract = Contract.objects.filter(
            employee_id__id=emp_id, contract_status="active"
        ).first()

        if not contract:
            continue

        if start_datetime is not None and start_datetime < contract.contract_start_date:
            errors.append(
                _(
                    "The %(employee)s's contract start date is smaller than pay period start date"
                )
                % {"employee": contract.employee_id}
            )
            valid = False

    if (
        start_datetime is not None
        and end_datetime is not None
        and start_datetime > end_datetime
    ):
        errors.append(
            _("The end date must be greater than or equal to the start date.")
        )
        valid = False

    if end_datetime is not None:
        if end_datetime > datetime.today().date():
            errors.append(_("The end date cannot be in the future."))
            valid = False

    return JsonResponse(
        {
            "valid": valid,
            "errors": errors,
        }
    )


@login_required
@permission_required("payroll.view_payslip")
def view_individual_payslip(request, employee_id, start_date, end_date):
    """
    This method is used to render the template for viewing a payslip.
    """

    payslip_data = payroll_calculation(employee_id, start_date, end_date)
    if not payslip_data:
        return HorillaRedirect(
            request,
            message=_(
                "Payslip data not found for the specified employee and date range."
            ),
        )
    return render(
        request,
        "payroll/payslip/individual_payslip.html",
        payslip_data,
    )


@login_required
@never_cache
def view_payslip(request):
    """
    This method is used to render the template for viewing a payslip.
    """
    if request.user.has_perm("payroll.view_payslip"):
        payslips = Payslip.objects.all()
    else:
        payslips = Payslip.objects.filter(employee_id__employee_user_id=request.user)
    export_column = forms.PayslipExportColumnForm()
    filter_form = PayslipFilter(request.GET, payslips)
    payslips = filter_form.qs
    bulk_form = forms.GeneratePayslipForm()
    field = request.GET.get("group_by")
    if field in Payslip.__dict__.keys():
        payslips = payslips.filter(group_name__isnull=False).order_by(field)
    payslips = paginator_qry(payslips, request.GET.get("page"))
    previous_data = request.GET.urlencode()
    data_dict = parse_qs(previous_data)
    get_key_instances(Payslip, data_dict)
    return render(
        request,
        "payroll/payslip/view_payslips.html",
        {
            "payslips": payslips,
            "f": filter_form,
            "export_column": export_column,
            "export_filter": PayslipFilter(request.GET),
            "bulk_form": bulk_form,
            "filter_dict": data_dict,
            "gp_fields": PayslipReGroup.fields,
        },
    )


@login_required
@hx_request_required
def filter_payslip(request):
    """
    Filter and retrieve a list of payslips based on the provided query parameters.
    """
    query_string = request.GET.urlencode()
    if request.user.has_perm("payroll.view_payslip"):
        payslips = PayslipFilter(request.GET).qs
    else:
        emp_request = request.GET.copy()
        employee = Employee.objects.filter(employee_user_id=request.user.id).first()
        employee_id = employee.id
        emp_request["employee_id"] = str(employee_id)
        payslips = PayslipFilter(emp_request).qs
    template = "payroll/payslip/payslip_table.html"
    view = request.GET.get("view")
    if view == "card":
        template = "payroll/payslip/group_payslips.html"
        payslips = payslips.filter(group_name__isnull=False).order_by("-group_name")
    payslips = sortby(request, payslips, "sortby")
    data_dict = []
    if not request.GET.get("dashboard"):
        data_dict = parse_qs(query_string)
        get_key_instances(Payslip, data_dict)
    if "status" in data_dict:
        status_list = data_dict["status"]
        if len(status_list) > 1:
            data_dict["status"] = [status_list[-1]]
    field = request.GET.get("field")
    if field != "" and field is not None:
        payslips = group_by_queryset(payslips, field, request.GET.get("page"), "page")
        template = "payroll/payslip/group_by.html"
    else:
        payslips = paginator_qry(payslips, request.GET.get("page"))
    return render(
        request,
        template,
        {
            "payslips": payslips,
            "pd": query_string,
            "filter_dict": data_dict,
        },
    )


@login_required
@permission_required("payroll.change_payslip")
def payslip_export(request):
    """
    This view exports payslip data based on selected fields and filters,
    and generates an Excel file for download.
    """
    if request.META.get("HTTP_HX_REQUEST"):
        return render(
            request,
            "payroll/payslip/payslip_export_filter.html",
            {
                "export_column": forms.PayslipExportColumnForm(),
                "export_filter": PayslipFilter(request.GET),
            },
        )

    choices_mapping = {
        "draft": _("Draft"),
        "review_ongoing": _("Review Ongoing"),
        "confirmed": _("Confirmed"),
        "paid": _("Paid"),
    }
    selected_columns = []
    payslips_data = {}
    payslips = PayslipFilter(request.GET).qs
    today_date = date.today().strftime("%Y-%m-%d")
    file_name = f"Payslip_excel_{today_date}.xlsx"
    selected_fields = request.GET.getlist("selected_fields")
    form = forms.PayslipExportColumnForm()

    if not selected_fields:
        selected_fields = form.fields["selected_fields"].initial
        ids = request.GET.get("ids", "[]")
        id_list = json.loads(ids)
        payslips = Payslip.objects.filter(id__in=id_list)

    for field in forms.excel_columns:
        value = field[0]
        key = field[1]
        if value in selected_fields:
            selected_columns.append((value, key))

    for column_value, column_name in selected_columns:
        nested_attributes = column_value.split("__")
        payslips_data[column_name] = []
        for payslip in payslips:
            value = payslip
            for attr in nested_attributes:
                value = getattr(value, attr, None)
                if value is None:
                    break
            data = str(value) if value is not None else ""
            if column_name == "Status":
                data = choices_mapping.get(value, "")

            if type(value) == date:
                date_format = request.user.employee_get.get_date_format()
                start_date = datetime.strptime(str(value), "%Y-%m-%d").date()

                for format_name, format_string in settings.HORILLA_DATE_FORMATS.items():
                    if format_name == date_format:
                        data = start_date.strftime(format_string)
            else:
                data = str(value) if value is not None else ""
            payslips_data[column_name].append(data)

    data_frame = pd.DataFrame(data=payslips_data)
    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = f'attachment; filename="{file_name}"'

    writer = pd.ExcelWriter(response, engine="xlsxwriter")
    data_frame.style.map(lambda x: "text-align: center").to_excel(
        writer, index=False, sheet_name="Sheet1"
    )
    worksheet = writer.sheets["Sheet1"]
    worksheet.set_column("A:Z", 20)
    writer.close()
    return response

def _parse_float(value):
    """Converte un valore stringa con separatore decimale virgola in float."""
    if value is None:
        return None
    s = str(value).strip()
    if s == "" or s == "nan":
        return None
    s = s.replace(".", "").replace(",", ".")
    try:
        return float(s)
    except (ValueError, TypeError):
        return None

def _parse_decimal(value):
    """Converte una stringa con separatori italiani (punto migliaia, virgola decimale) in Decimal."""
    from decimal import Decimal, InvalidOperation
    if value is None:
        return None
    s = str(value).strip()
    if s == "" or s == "nan":
        return None
    s = s.replace(".", "").replace(",", ".")
    try:
        return Decimal(s)
    except (InvalidOperation, ValueError):
        return None

def _parse_date_it(value):
    """Converte una stringa data italiana dd.mm.yyyy o dd/mm/yyyy in oggetto date."""
    from datetime import datetime as _dt
    if value is None:
        return None
    s = str(value).strip()
    if not s:
        return None
    for fmt in ("%d.%m.%Y", "%d/%m/%Y", "%Y-%m-%d"):
        try:
            return _dt.strptime(s, fmt).date()
        except ValueError:
            continue
    return None

@login_required
def import_payslip_presenze(request):
    """
    GET  – mostra il form per selezionare mese/anno e caricare il CSV.
    POST – valida il file e importa i dati nella tabella payslip_presenze.
    """
    months = [
        (1, "Gennaio"), (2, "Febbraio"), (3, "Marzo"), (4, "Aprile"),
        (5, "Maggio"), (6, "Giugno"), (7, "Luglio"), (8, "Agosto"),
        (9, "Settembre"), (10, "Ottobre"), (11, "Novembre"), (12, "Dicembre"),
    ]
    current_year = date.today().year

    def _render_pres(extra=None):
        ctx = {"months": months, "current_year": current_year}
        if extra:
            ctx.update(extra)
        return render(request, "payroll/payslip/import_payslip_presenze.html", ctx)

    if request.method == "GET":
        request.session.pop('import_presenze_pending', None)
        return _render_pres()

    # ── Step 2b: conflitti – l'utente ha scelto replace/keep ──
    if request.POST.get("step") == "conflicts":
        pending = request.session.get('import_presenze_pending')
        if not pending:
            messages.error(request, _("Sessione scaduta. Ricaricare il file."))
            return _render_pres()
        mese, anno = pending['mese'], pending['anno']
        action = request.POST.get('conflict_action', 'keep')
        conflict_matricole = set(pending.get('conflict_matricole', []))
        if action == 'replace':
            PayslipPresenze.objects.filter(
                mese=mese, anno=anno, matricola__in=conflict_matricole
            ).delete()
        else:
            pending['objects_data'] = [
                od for od in pending['objects_data']
                if od.get('matricola') not in conflict_matricole
            ]
        request.session['import_presenze_pending'] = pending
        # Se ci sono ancora matricole non trovate, vai a step 2
        if pending.get('unmatched'):
            return _render_pres({
                "mese": mese, "anno": anno,
                "step2": True,
                "unmatched": pending['unmatched'],
                "n_total": len(pending['objects_data']),
                "n_unmatched": len(pending['unmatched']),
            })
        # Altrimenti crea direttamente
        from datetime import date as _date_cls
        objects_to_create = []
        for od in pending['objects_data']:
            kwargs = dict(od)
            if kwargs.get('data_ass'):
                kwargs['data_ass'] = _date_cls.fromisoformat(kwargs['data_ass'])
            objects_to_create.append(PayslipPresenze(**kwargs))
        PayslipPresenze.objects.bulk_create(objects_to_create)
        request.session.pop('import_presenze_pending', None)
        messages.success(request, _("Importate {} righe per {:02d}/{}.").format(len(objects_to_create), mese, anno))
        return redirect("view-payslip")

    # ── Step 2: l'utente ha fornito i badge_id per le matricole non trovate ──
    if request.POST.get("step") == "2":
        pending = request.session.get('import_presenze_pending')
        if not pending:
            messages.error(request, _("Sessione scaduta. Ricaricare il file."))
            return render(request, "payroll/payslip/import_payslip_presenze.html",
                          {"months": months, "current_year": current_year})
        mese, anno = pending['mese'], pending['anno']
        # matricola → badge_id (cod_dip) inserito dall'utente
        override_map = {
            key[4:]: val.strip()
            for key, val in request.POST.items()
            if key.startswith('bid_') and val.strip()
        }
        from datetime import date as _date_cls
        objects_to_create = []
        for od in pending['objects_data']:
            kwargs = dict(od)
            if kwargs.get('data_ass'):
                kwargs['data_ass'] = _date_cls.fromisoformat(kwargs['data_ass'])
            mat = kwargs.get('matricola')
            if kwargs.get('cod_dip') is None and mat and mat in override_map:
                kwargs['cod_dip'] = override_map[mat]
            objects_to_create.append(PayslipPresenze(**kwargs))
        PayslipPresenze.objects.bulk_create(objects_to_create)
        request.session.pop('import_presenze_pending', None)
        n_senza = sum(1 for o in objects_to_create if o.cod_dip is None and o.matricola)
        msg = _("Importate {} righe per {:02d}/{}.").format(len(objects_to_create), mese, anno)
        if n_senza:
            msg += " " + _("{} righe senza cod_dip (matricola non risolta).").format(n_senza)
        messages.success(request, msg)
        return redirect("view-payslip")

    mese_str = request.POST.get("mese", "").strip()
    anno_str = request.POST.get("anno", "").strip()
    csv_file = request.FILES.get("file")

    # --- validazione parametri base ---
    if not mese_str or not anno_str:
        messages.error(request, _("Selezionare mese e anno prima di importare."))
        return render(request, "payroll/payslip/import_payslip_presenze.html",
                      {"mese": mese_str, "anno": anno_str, "months": months, "current_year": current_year})

    try:
        mese = int(mese_str)
        anno = int(anno_str)
        if not (1 <= mese <= 12):
            raise ValueError
    except ValueError:
        messages.error(request, _("Mese o anno non validi."))
        return render(request, "payroll/payslip/import_payslip_presenze.html",
                      {"mese": mese_str, "anno": anno_str, "months": months, "current_year": current_year})

    if not csv_file:
        messages.error(request, _("Nessun file caricato."))
        return _render_pres({"mese": mese, "anno": anno})

    # --- lettura CSV ---
    import csv as csv_module
    import io
    from datetime import datetime as dt_parse

    try:
        content = csv_file.read().decode("utf-8", errors="replace")
        reader = csv_module.DictReader(io.StringIO(content), delimiter=";")
        rows = list(reader)
    except Exception as exc:
        messages.error(request, _("Impossibile leggere il file CSV: {}").format(exc))
        return render(request, "payroll/payslip/import_payslip_presenze.html",
                      {"mese": mese, "anno": anno, "months": months, "current_year": current_year})

    if not rows:
        messages.error(request, _("Il file CSV è vuoto."))
        return render(request, "payroll/payslip/import_payslip_presenze.html",
                      {"mese": mese, "anno": anno, "months": months, "current_year": current_year})

    # --- validazione mese/anno su tutte le righe ---
    invalid_rows = []       # (n_riga, trovato_mese, trovato_anno)
    found_combos = set()    # coppie mese/anno trovate nel file
    for idx, row in enumerate(rows, start=2):  # riga 1 = header
        try:
            row_mese = int(str(row.get("mese", "")).strip())
            row_anno = int(str(row.get("anno", "")).strip())
        except (ValueError, TypeError):
            invalid_rows.append((idx, "?", "?"))
            continue
        found_combos.add((row_mese, row_anno))
        if row_mese != mese or row_anno != anno:
            invalid_rows.append((idx, row_mese, row_anno))

    if invalid_rows:
        found_str = ", ".join(
            f"{m:02d}/{a}" for m, a in sorted(found_combos) if (m, a) != (mese, anno)
        )
        rows_str = ", ".join(str(r[0]) for r in invalid_rows[:20])
        suffix = f" ... e altre {len(invalid_rows) - 20}" if len(invalid_rows) > 20 else ""
        messages.error(
            request,
            _(
                "Hai selezionato {sel_mese:02d}/{sel_anno}, ma il file CSV contiene "
                "righe con mese/anno: {found}. "
                "Righe non conformi: {rows}{suffix}."
            ).format(
                sel_mese=mese,
                sel_anno=anno,
                found=found_str or "non leggibile",
                rows=rows_str,
                suffix=suffix,
            ),
        )
        return render(request, "payroll/payslip/import_payslip_presenze.html",
                      {"mese": mese, "anno": anno, "months": months, "current_year": current_year})

    # --- importazione ---
    day_fields = [str(i) for i in range(1, 32)]
    objects_data = []
    seen_unmatched = set()
    unmatched_list = []  # [{matricola, lavoratore}]

    for row in rows:
        # data_ass: dd/mm/yyyy -> yyyy-mm-dd
        data_ass = None
        raw_data_ass = str(row.get("data ass", "")).strip()
        if raw_data_ass:
            try:
                data_ass = dt_parse.strptime(raw_data_ass, "%d/%m/%Y").date()
            except ValueError:
                try:
                    data_ass = dt_parse.strptime(raw_data_ass, "%Y-%m-%d").date()
                except ValueError:
                    data_ass = None

        day_values = {}
        for i, col in enumerate(day_fields, start=1):
            day_values[f"day_{i}"] = _parse_float(row.get(col))

        mat = str(row.get("matricola", "")).strip()
        lavoratore_val = str(row.get("lavoratore", "")).strip() or None
        mat_mese_anno = f"{mat}_{mese:02d}_{anno}" if mat else None

        # Ricava badge_id dall'anagrafica dipendenti tramite codice_paghe == matricola
        cod_dip = None
        if mat:
            emp = Employee.objects.filter(codice_paghe=mat).first()
            if emp:
                cod_dip = emp.badge_id
            elif mat not in seen_unmatched:
                seen_unmatched.add(mat)
                unmatched_list.append({'matricola': mat, 'lavoratore': lavoratore_val or mat})

        row_data = {
            'dl':               str(row.get("dl", "")).strip() or None,
            'fil':              str(row.get("fil", "")).strip() or None,
            'cc':               str(row.get("cc", "")).strip() or None,
            'rag_soc':          str(row.get("rag.soc.", "")).strip() or None,
            'matricola':        mat or None,
            'lavoratore':       lavoratore_val,
            'qp':               str(row.get("qp", "")).strip() or None,
            'data_ass':         data_ass.isoformat() if data_ass else None,
            'livello':          str(row.get("livello", "")).strip() or None,
            'desc_liv':         str(row.get(" desc.liv.", "")).strip() or None,
            'pt':               str(row.get("pt", "")).strip() or None,
            'perc_pt':          _parse_float(row.get("%pt")),
            'perc_turn':        _parse_float(row.get("%turn")),
            'mese':             mese,
            'anno':             anno,
            'matricola_mese_anno': mat_mese_anno,
            'cod_voce':         int(str(row.get("cod.voce", "")).strip()) if str(row.get("cod.voce", "")).strip() else None,
            'desc_voce':        str(row.get("desc.voce", "")).strip() or None,
            'aliq_voce':        _parse_float(row.get("aliq.voce")),
            'ore_tot':          _parse_float(row.get("ore tot")),
            'gg_tot':           _parse_float(row.get("gg tot")),
            'periodo_elab':     str(row.get("periodo elab.", "")).strip() or None,
            'cod_dip':          cod_dip,
            **day_values,
        }
        objects_data.append(row_data)

    # --- Controllo conflitti: matricole già presenti per mese/anno ---
    matricole_nel_file = {od['matricola'] for od in objects_data if od.get('matricola')}
    gia_presenti_pres = set(
        PayslipPresenze.objects.filter(mese=mese, anno=anno, matricola__in=matricole_nel_file)
        .values_list('matricola', flat=True).distinct()
    )
    if gia_presenti_pres:
        request.session['import_presenze_pending'] = {
            'mese': mese, 'anno': anno,
            'objects_data': objects_data,
            'unmatched': unmatched_list,
            'conflict_matricole': sorted(gia_presenti_pres),
        }
        return _render_pres({
            'mese': mese, 'anno': anno,
            'warn_conflicts': True,
            'conflict_matricole': sorted(gia_presenti_pres),
            'n_conflicts': len(gia_presenti_pres),
            'n_total_matricole': len(matricole_nel_file),
        })

    # Matricole senza badge_id: chiedi all'utente
    if unmatched_list:
        request.session['import_presenze_pending'] = {
            'mese': mese,
            'anno': anno,
            'objects_data': objects_data,
            'unmatched': unmatched_list,
        }
        return _render_pres({
            "mese": mese, "anno": anno,
            "step2": True,
            "unmatched": unmatched_list,
            "n_total": len(objects_data),
            "n_unmatched": len(unmatched_list),
        })

    # Nessun non-trovato: crea direttamente
    from datetime import date as _date_cls
    objects_to_create = []
    for od in objects_data:
        kwargs = dict(od)
        if kwargs.get('data_ass'):
            kwargs['data_ass'] = _date_cls.fromisoformat(kwargs['data_ass'])
        objects_to_create.append(PayslipPresenze(**kwargs))

    PayslipPresenze.objects.bulk_create(objects_to_create)
    messages.success(
        request,
        _("Importate {} righe per {:02d}/{}.").format(len(objects_to_create), mese, anno),
    )
    return redirect("view-payslip")

@login_required
def delete_payslip_presenze(request):
    """
    POST – cancella tutte le righe di payslip_presenze per il mese/anno indicati.
    GET  – mostra il form di conferma cancellazione.
    """
    months = [
        (1, "Gennaio"), (2, "Febbraio"), (3, "Marzo"), (4, "Aprile"),
        (5, "Maggio"), (6, "Giugno"), (7, "Luglio"), (8, "Agosto"),
        (9, "Settembre"), (10, "Ottobre"), (11, "Novembre"), (12, "Dicembre"),
    ]
    current_year = date.today().year

    if request.method == "GET":
        return render(request, "payroll/payslip/delete_payslip_presenze.html",
                      {"months": months, "current_year": current_year})

    mese_str = request.POST.get("mese", "").strip()
    anno_str = request.POST.get("anno", "").strip()

    try:
        mese = int(mese_str)
        anno = int(anno_str)
        if not (1 <= mese <= 12):
            raise ValueError
    except (ValueError, TypeError):
        messages.error(request, _("Mese o anno non validi."))
        return render(request, "payroll/payslip/delete_payslip_presenze.html",
                      {"mese": mese_str, "anno": anno_str, "months": months, "current_year": current_year})

    deleted_count, _del = PayslipPresenze.objects.filter(mese=mese, anno=anno).delete()
    messages.success(
        request,
        _("Eliminati {} record per {:02d}/{}.").format(deleted_count, mese, anno),
    )
    return redirect("view-payslip")

@login_required
def delete_payslip_corpo(request):
    """
    POST – cancella tutte le righe di payslip_corpo per il mese/anno indicati.
    GET  – mostra il form di conferma cancellazione.
    """
    from payroll.models.models import PayslipCorpo

    months = [
        (1, "Gennaio"), (2, "Febbraio"), (3, "Marzo"), (4, "Aprile"),
        (5, "Maggio"), (6, "Giugno"), (7, "Luglio"), (8, "Agosto"),
        (9, "Settembre"), (10, "Ottobre"), (11, "Novembre"), (12, "Dicembre"),
    ]
    current_year = date.today().year

    if request.method == "GET":
        return render(request, "payroll/payslip/delete_payslip_corpo.html",
                      {"months": months, "current_year": current_year})

    mese_str = request.POST.get("mese", "").strip()
    anno_str = request.POST.get("anno", "").strip()

    try:
        mese = int(mese_str)
        anno = int(anno_str)
        if not (1 <= mese <= 12):
            raise ValueError
    except (ValueError, TypeError):
        messages.error(request, _("Mese o anno non validi."))
        return render(request, "payroll/payslip/delete_payslip_corpo.html",
                      {"mese": mese_str, "anno": anno_str, "months": months, "current_year": current_year})

    from django.db import transaction
    from payroll.models.models import Payslip

    with transaction.atomic():
        # Trova i Payslip collegati tramite FK prima di cancellare le righe corpo
        payslip_ids = list(
            PayslipCorpo.objects.filter(mese=mese, anno=anno)
            .exclude(payslip_id=None)
            .values_list("payslip_id", flat=True)
            .distinct()
        )
        deleted_count, _del_detail = PayslipCorpo.objects.filter(mese=mese, anno=anno).delete()
        payslip_deleted = 0
        if payslip_ids:
            payslip_deleted, _ps_detail = Payslip.objects.filter(id__in=payslip_ids).delete()

    msg = _("Eliminati {} record corpo per {:02d}/{}.").format(deleted_count, mese, anno)
    if payslip_deleted:
        msg += " " + _("Buste paga eliminate: {}.").format(payslip_deleted)
    messages.success(request, msg)
    return redirect("view-payslip")

@login_required
def import_payslip_corpo(request):
    """
    GET  – mostra il form per selezionare mese/anno e caricare il CSV del corpo busta.
    POST – importa il file CSV nella tabella payslip_corpo.
    Mese e anno NON sono presenti nel CSV: vengono impostati dall'utente nel form
    e scritti su ogni riga importata.
    """
    from payroll.models.models import PayslipCorpo

    months = [
        (1, "Gennaio"), (2, "Febbraio"), (3, "Marzo"), (4, "Aprile"),
        (5, "Maggio"), (6, "Giugno"), (7, "Luglio"), (8, "Agosto"),
        (9, "Settembre"), (10, "Ottobre"), (11, "Novembre"), (12, "Dicembre"),
    ]
    current_year = date.today().year

    def _render_corpo(extra=None):
        ctx = {"months": months, "current_year": current_year}
        if extra:
            ctx.update(extra)
        return render(request, "payroll/payslip/import_payslip_corpo.html", ctx)

    if request.method == "GET":
        request.session.pop('import_corpo_pending', None)
        return _render_corpo()

    # ── Step conflicts: l'utente ha scelto cosa fare con i conflitti ──
    if request.POST.get("step") == "conflicts":
        pending = request.session.get('import_corpo_pending')
        if not pending:
            messages.error(request, _("Sessione scaduta. Ricaricare il file."))
            return _render_corpo()
        mese, anno = pending['mese'], pending['anno']
        action = request.POST.get('conflict_action', 'keep')  # 'replace' o 'keep'
        objects_to_create = []
        for od in pending['objects_data']:
            obj = PayslipCorpo(**{k: v for k, v in od.items() if k not in ('assunzione', 'anzianita', 'data_pos')})
            from datetime import date as _dc
            obj.assunzione = _dc.fromisoformat(od['assunzione']) if od.get('assunzione') else None
            obj.anzianita  = _dc.fromisoformat(od['anzianita'])  if od.get('anzianita')  else None
            obj.data_pos   = _dc.fromisoformat(od['data_pos'])   if od.get('data_pos')   else None
            objects_to_create.append(obj)
        conflict_matricole = set(pending['conflict_matricole'])
        if action == 'replace':
            PayslipCorpo.objects.filter(mese=mese, anno=anno, matricola__in=conflict_matricole).delete()
        else:
            objects_to_create = [o for o in objects_to_create if o.matricola not in conflict_matricole]
        request.session.pop('import_corpo_pending', None)
        # prosegui con la bulk_create e sincronizzazione
        # (il codice sotto viene riutilizzato)
        _corpo_bulk_and_sync(request, mese, anno, objects_to_create, skipped=pending.get('skipped', 0))
        return redirect("view-payslip")

    mese_str = request.POST.get("mese", "").strip()
    anno_str = request.POST.get("anno", "").strip()
    csv_file = request.FILES.get("file")

    if not mese_str or not anno_str:
        messages.error(request, _("Selezionare mese e anno prima di importare."))
        return render(request, "payroll/payslip/import_payslip_corpo.html",
                      {"mese": mese_str, "anno": anno_str, "months": months, "current_year": current_year})

    try:
        mese = int(mese_str)
        anno = int(anno_str)
        if not (1 <= mese <= 12):
            raise ValueError
    except ValueError:
        messages.error(request, _("Mese o anno non validi."))
        return render(request, "payroll/payslip/import_payslip_corpo.html",
                      {"mese": mese_str, "anno": anno_str, "months": months, "current_year": current_year})

    if not csv_file:
        messages.error(request, _("Nessun file caricato."))
        return render(request, "payroll/payslip/import_payslip_corpo.html",
                      {"mese": mese, "anno": anno, "months": months, "current_year": current_year})

    import csv as csv_module
    import io

    try:
        content = csv_file.read().decode("utf-8", errors="replace")
        reader = csv_module.reader(io.StringIO(content), delimiter=";")
        all_rows = list(reader)
    except Exception as exc:
        messages.error(request, _("Impossibile leggere il file CSV: {}").format(exc))
        return render(request, "payroll/payslip/import_payslip_corpo.html",
                      {"mese": mese, "anno": anno, "months": months, "current_year": current_year})

    if len(all_rows) < 2:
        messages.error(request, _("Il file CSV è vuoto o privo di dati."))
        return render(request, "payroll/payslip/import_payslip_corpo.html",
                      {"mese": mese, "anno": anno, "months": months, "current_year": current_year})

    # colonne attese (indice 0-based), ignoriamo l'ultima colonna vuota (trailing ;)
    # 0:Codice DL, 1:Denominazione, 2:Filiale, 3:C.Costo, 4:Reparto,
    # 5:Matricola, 6:Cognome, 7:Nome, 8:qp, 9:ass., 10:anz.,
    # 11:cod.pos., 12:data pos., 13:liq., 14:Cod.voce, 15:Descrizione.,
    # 16:aliq./%lav., 17:h/g/n /%d.l., 18:Dato base/imponibile,
    # 19:Importo/ctr lav, 20:d.b.tfr, 21:Imp.tfr/ctr.dl
    EXPECTED_COLS = 22

    objects_to_create = []
    skipped = 0

    for idx, row in enumerate(all_rows[1:], start=2):  # salta header
        if len(row) < EXPECTED_COLS:
            skipped += 1
            continue

        def _cell(i):
            return row[i].strip() if i < len(row) else ""

        def _int_or_none(i):
            v = _cell(i)
            try:
                return int(v) if v else None
            except ValueError:
                return None

        obj = PayslipCorpo(
            mese=mese,
            anno=anno,
            codice_dl=_cell(0)[:20] or None,
            denominazione=_cell(1)[:100] or None,
            filiale=_cell(2)[:20] or None,
            c_costo=_int_or_none(3),
            reparto=_int_or_none(4),
            matricola=_cell(5)[:20] or None,
            cognome=_cell(6)[:100] or None,
            nome=_cell(7)[:100] or None,
            qp=_cell(8)[:10] or None,
            assunzione=_parse_date_it(_cell(9)),
            anzianita=_parse_date_it(_cell(10)),
            cod_pos=_int_or_none(11),
            data_pos=_parse_date_it(_cell(12)),
            liq=_cell(13)[:20] or None,
            cod_voce=_int_or_none(14),
            descrizione_voce=_cell(15)[:100] or None,
            aliq_perc_lav=_parse_decimal(_cell(16)),
            unita=_parse_decimal(_cell(17)),
            dato_base_imponibile=_parse_decimal(_cell(18)),
            importo_ctr_lav=_parse_decimal(_cell(19)),
            db_tfr=_parse_decimal(_cell(20)),
            imp_tfr_ctr_dl=_parse_decimal(_cell(21)),
        )
        objects_to_create.append(obj)

    if not objects_to_create:
        messages.error(request, _("Nessuna riga valida trovata nel file CSV."))
        return _render_corpo({"mese": mese, "anno": anno})

    # --- Controllo conflitti: matricole già presenti per mese/anno ---
    matricole_nel_file = {o.matricola for o in objects_to_create if o.matricola}
    gia_presenti = set(
        PayslipCorpo.objects.filter(mese=mese, anno=anno, matricola__in=matricole_nel_file)
        .values_list('matricola', flat=True).distinct()
    )
    if gia_presenti:
        # Serializza date come stringhe ISO per la sessione
        def _ser(obj):
            d = obj.__dict__.copy()
            d.pop('_state', None)
            d.pop('id', None)
            d.pop('payslip_id', None)
            for fld in ('assunzione', 'anzianita', 'data_pos'):
                if d.get(fld) is not None:
                    d[fld] = d[fld].isoformat()
            for fld in ('aliq_perc_lav', 'unita', 'dato_base_imponibile',
                         'importo_ctr_lav', 'db_tfr', 'imp_tfr_ctr_dl'):
                if d.get(fld) is not None:
                    d[fld] = str(d[fld])
            return d
        request.session['import_corpo_pending'] = {
            'mese': mese, 'anno': anno,
            'objects_data': [_ser(o) for o in objects_to_create],
            'conflict_matricole': sorted(gia_presenti),
            'skipped': skipped,
        }
        return _render_corpo({
            'mese': mese, 'anno': anno,
            'warn_conflicts': True,
            'conflict_matricole': sorted(gia_presenti),
            'n_conflicts': len(gia_presenti),
            'n_total_matricole': len(matricole_nel_file),
        })

    _corpo_bulk_and_sync(request, mese, anno, objects_to_create, skipped)
    return redirect("view-payslip")

def _corpo_bulk_and_sync(request, mese, anno, objects_to_create, skipped=0):
    """Bulk-crea le righe corpo e sincronizza Payslip e salary_hour."""
    from payroll.models.models import PayslipCorpo, Payslip
    from decimal import Decimal

    PayslipCorpo.objects.bulk_create(objects_to_create)

    # --- Sincronizzazione payroll_payslip e update salary_hour ---
    import calendar
    from django.db import transaction
    from employee.models import Employee as Emp, EmployeeWorkInformation
    from payroll.models.models import Payslip

    start_date = date(anno, mese, 1)
    end_date = date(anno, mese, calendar.monthrange(anno, mese)[1])

    # Raggruppa le righe per matricola sommando cod_voce=852 e cod_voce=800
    net_pay_by_matricola = {}
    for obj in objects_to_create:
        if obj.cod_voce in (852, 800) and obj.matricola and obj.importo_ctr_lav is not None:
            net_pay_by_matricola[obj.matricola] = (
                net_pay_by_matricola.get(obj.matricola, 0.0) + float(obj.importo_ctr_lav)
            )

    salary_updated = 0
    salary_not_found = []
    payslip_created = 0
    payslip_updated = 0
    payslip_no_852 = []
    payslip_no_emp = []

    # Raccogli tutte le matricole uniche presenti nel CSV
    all_matricole = {obj.matricola for obj in objects_to_create if obj.matricola}

    with transaction.atomic():
        for matricola in all_matricole:
            # --- update salary_hour per cod_voce=300 ---
            voce_300 = next(
                (o for o in objects_to_create if o.matricola == matricola and o.cod_voce == 300 and o.dato_base_imponibile is not None),
                None,
            )
            emp = Emp.objects.filter(codice_paghe=matricola).first()
            if not emp:
                salary_not_found.append(matricola)
                if matricola not in net_pay_by_matricola:
                    payslip_no_emp.append(matricola)
                continue

            if voce_300:
                updated = EmployeeWorkInformation.objects.filter(employee_id=emp).update(
                    salary_hour=voce_300.dato_base_imponibile
                )
                if updated:
                    salary_updated += 1

            # --- Crea/aggiorna Payslip da cod_voce=852 ---
            if matricola not in net_pay_by_matricola:
                payslip_no_852.append(matricola)
                continue

            net_pay = net_pay_by_matricola[matricola]
            payslip, created = Payslip.objects.get_or_create(
                employee_id=emp,
                start_date=start_date,
                end_date=end_date,
                defaults={
                    "net_pay": net_pay,
                    "gross_pay": net_pay,
                    "basic_pay": net_pay,
                    "deduction": 0,
                    "contract_wage": net_pay,
                    "status": "paid",
                    "pay_head_data": {"source": "import_corpo", "cod_voce_852_800": net_pay},
                    "sent_to_employee": False,
                },
            )
            if not created:
                payslip.net_pay = net_pay
                payslip.gross_pay = net_pay
                payslip.status = "paid"
                payslip.pay_head_data = {"source": "import_corpo", "cod_voce_852_800": net_pay}
                payslip.save(update_fields=["net_pay", "gross_pay", "status", "pay_head_data"])
                payslip_updated += 1
            else:
                payslip_created += 1

            # Collega tutte le righe corpo di questa matricola/mese/anno al Payslip
            PayslipCorpo.objects.filter(
                matricola=matricola, mese=mese, anno=anno
            ).update(payslip=payslip)

    msg = _("Importate {} righe per {:02d}/{}.").format(len(objects_to_create), mese, anno)
    if skipped:
        msg += " " + _("{} righe ignorate (colonne insufficienti).").format(skipped)
    if salary_updated:
        msg += " " + _("salary_hour aggiornato per {} dipendenti.").format(salary_updated)
    if payslip_created:
        msg += " " + _("Buste paga create: {}.").format(payslip_created)
    if payslip_updated:
        msg += " " + _("Buste paga aggiornate: {}.").format(payslip_updated)
    if salary_not_found:
        not_found_str = ", ".join(sorted(set(salary_not_found)))
        msg += " " + _("Matricole non trovate in anagrafica: {}.").format(not_found_str)
    if payslip_no_852:
        msg += " " + _("Nessun cod_voce 852/800 (netto) per: {}.").format(", ".join(sorted(set(payslip_no_852))))
    messages.success(request, msg)

@login_required
def import_payslip_importi(request):
    """
    GET  – mostra il form per selezionare mese/anno e caricare il file Excel.
    POST – importa i dati dalla tabella payslip_importi.

    Struttura Excel attesa (senza intestazioni significative, prima riga skippata):
      col 0 = neg
      col 1 = (ignorata)
      col 2 = badge_id
      col 3 = importo
    La colonna matricola viene ricavata da employee_employee.codice_paghe
    dove badge_id corrisponde a quello della riga in importazione.
    """
    from payroll.models.models import PayslipImporti

    current_year = date.today().year

    def _render(extra=None):
        ctx = {"months": _MONTHS_IT, "current_year": current_year}
        if extra:
            ctx.update(extra)
        return render(request, "payroll/payslip/import_payslip_importi.html", ctx)

    if request.method == "GET":
        request.session.pop('import_importi_pending', None)
        return _render()

    # ── Step 2b: conflitti badge_id – l'utente ha scelto replace/keep ──
    if request.POST.get("step") == "conflicts":
        pending = request.session.get('import_importi_pending')
        if not pending:
            messages.error(request, _("Sessione scaduta. Ricaricare il file."))
            return _render()
        mese, anno = pending['mese'], pending['anno']
        action = request.POST.get('conflict_action', 'keep')
        conflict_badge_ids = set(pending.get('conflict_badge_ids', []))
        if action == 'replace':
            PayslipImporti.objects.filter(
                mese=mese, anno=anno, badge_id__in=conflict_badge_ids
            ).delete()
        else:
            pending['objects_data'] = [
                od for od in pending['objects_data']
                if od.get('badge_id') not in conflict_badge_ids
            ]
        request.session['import_importi_pending'] = pending
        # Se ci sono ancora badge_id non trovati in anagrafica, vai a step 2
        if pending.get('unmatched'):
            return _render({
                'mese': mese, 'anno': anno,
                'step2': True,
                'unmatched': pending['unmatched'],
                'n_total': len(pending['objects_data']),
                'n_unmatched': len(pending['unmatched']),
            })
        # Altrimenti crea direttamente
        from decimal import Decimal
        if pending.get('needs_delete'):
            PayslipImporti.objects.filter(mese=mese, anno=anno).delete()
        PayslipImporti.objects.bulk_create([
            PayslipImporti(
                mese=mese, anno=anno,
                neg=od.get('neg'), badge_id=od.get('badge_id'),
                matricola=od.get('matricola'),
                importo=Decimal(od['importo']) if od.get('importo') else None,
            )
            for od in pending['objects_data']
        ])
        request.session.pop('import_importi_pending', None)
        messages.success(request, _("Importate {} righe per {:02d}/{}.").format(len(pending['objects_data']), mese, anno))
        return redirect("view-payslip")

    # ── Step 2: l'utente ha fornito le matricole per i badge_id non trovati ──
    if request.POST.get("step") == "2":
        from decimal import Decimal
        pending = request.session.get('import_importi_pending')
        if not pending:
            messages.error(request, _("Sessione scaduta. Ricaricare il file."))
            return _render()
        mese, anno = pending['mese'], pending['anno']
        override_map = {
            key[4:]: val.strip()
            for key, val in request.POST.items()
            if key.startswith('mat_') and val.strip()
        }
        to_create = []
        for od in pending['objects_data']:
            badge_id = od.get('badge_id')
            matricola = od.get('matricola') or override_map.get(badge_id or '')
            importo = Decimal(od['importo']) if od.get('importo') else None
            to_create.append(PayslipImporti(
                mese=mese, anno=anno,
                neg=od.get('neg'), badge_id=badge_id,
                matricola=matricola or None, importo=importo,
            ))
        if pending.get('needs_delete'):
            PayslipImporti.objects.filter(mese=mese, anno=anno).delete()
        PayslipImporti.objects.bulk_create(to_create)
        request.session.pop('import_importi_pending', None)
        n_senza = sum(1 for o in to_create if o.matricola is None)
        msg = _("Importate {} righe per {:02d}/{}.").format(len(to_create), mese, anno)
        if n_senza:
            msg += " " + _("{} senza matricola (badge_id non risolto).").format(n_senza)
        messages.success(request, msg)
        return redirect("view-payslip")

    mese_str = request.POST.get("mese", "").strip()
    anno_str = request.POST.get("anno", "").strip()
    excel_file = request.FILES.get("file")
    confirm_overwrite = request.POST.get("confirm_overwrite", "")

    # --- validazione mese/anno ---
    if not mese_str or not anno_str:
        messages.error(request, _("Selezionare mese e anno prima di importare."))
        return _render({"mese": mese_str, "anno": anno_str})

    try:
        mese = int(mese_str)
        anno = int(anno_str)
        if not (1 <= mese <= 12):
            raise ValueError
    except (ValueError, TypeError):
        messages.error(request, _("Mese o anno non validi."))
        return _render({"mese": mese_str, "anno": anno_str})

    if not excel_file:
        messages.error(request, _("Nessun file caricato."))
        return _render({"mese": mese, "anno": anno})

    # --- lettura Excel con pandas ---
    try:
        import io
        df = pd.read_excel(io.BytesIO(excel_file.read()), header=None, dtype=str)
    except Exception as exc:
        messages.error(request, _("Impossibile leggere il file Excel: {}").format(exc))
        return _render({"mese": mese, "anno": anno})

    # Salta la prima riga (intestazione)
    df = df.iloc[1:].reset_index(drop=True)

    if df.empty:
        messages.error(request, _("Il file Excel non contiene dati dopo l'intestazione."))
        return _render({"mese": mese, "anno": anno})

    # --- controlla se esistono già righe per mese/anno ---
    existing_count = PayslipImporti.objects.filter(mese=mese, anno=anno).count()
    if existing_count and confirm_overwrite != "yes":
        return _render({
            "mese": mese,
            "anno": anno,
            "existing_count": existing_count,
            "warn_existing": True,
        })

    needs_delete = existing_count > 0  # deferisce il delete al momento del bulk_create

    # --- costruzione oggetti da importare ---
    objects_data = []
    skipped = 0
    seen_unmatched = set()
    unmatched_list = []

    for idx, row in df.iterrows():
        cols = list(row)
        if len(cols) < 4:
            skipped += 1
            continue

        def _cell(i):
            v = cols[i]
            if v is None:
                return ""
            s = str(v).strip()
            return "" if s.lower() in ("nan", "none", "") else s

        neg_val = _cell(0) or None
        badge_id_val = _cell(2) or None
        importo_val = _parse_decimal(_cell(3))

        # ricava matricola da employee_employee tramite badge_id
        matricola_val = None
        if badge_id_val:
            emp = Employee.objects.filter(badge_id=badge_id_val).first()
            if emp and emp.codice_paghe:
                matricola_val = emp.codice_paghe
            elif badge_id_val not in seen_unmatched:
                seen_unmatched.add(badge_id_val)
                unmatched_list.append(badge_id_val)

        objects_data.append({
            'neg': neg_val,
            'badge_id': badge_id_val,
            'matricola': matricola_val,
            'importo': str(importo_val) if importo_val is not None else None,
        })

    if not objects_data:
        messages.error(request, _("Nessuna riga valida trovata nel file Excel."))
        return _render({"mese": mese, "anno": anno})

    # --- Controllo conflitti: badge_id già presenti per mese/anno ---
    badge_ids_nel_file = {od['badge_id'] for od in objects_data if od.get('badge_id')}
    gia_presenti_imp = set(
        PayslipImporti.objects.filter(mese=mese, anno=anno, badge_id__in=badge_ids_nel_file)
        .values_list('badge_id', flat=True).distinct()
    )
    if gia_presenti_imp:
        request.session['import_importi_pending'] = {
            'mese': mese, 'anno': anno,
            'needs_delete': needs_delete,
            'objects_data': objects_data,
            'unmatched': unmatched_list,
            'conflict_badge_ids': sorted(gia_presenti_imp),
        }
        return _render({
            'mese': mese, 'anno': anno,
            'warn_conflicts': True,
            'conflict_badge_ids': sorted(gia_presenti_imp),
            'n_conflicts': len(gia_presenti_imp),
            'n_total_badge': len(badge_ids_nel_file),
        })

    # Badge_id senza matricola: chiedi all'utente
    if unmatched_list:
        request.session['import_importi_pending'] = {
            'mese': mese,
            'anno': anno,
            'needs_delete': needs_delete,
            'objects_data': objects_data,
            'unmatched': unmatched_list,
        }
        return _render({
            'mese': mese, 'anno': anno,
            'step2': True,
            'unmatched': unmatched_list,
            'n_total': len(objects_data),
            'n_unmatched': len(unmatched_list),
            'skipped': skipped,
        })

    # Nessun non-trovato: crea direttamente
    from decimal import Decimal
    if needs_delete:
        PayslipImporti.objects.filter(mese=mese, anno=anno).delete()
    PayslipImporti.objects.bulk_create([
        PayslipImporti(
            mese=mese, anno=anno,
            neg=od['neg'], badge_id=od['badge_id'],
            matricola=od['matricola'],
            importo=Decimal(od['importo']) if od.get('importo') else None,
        )
        for od in objects_data
    ])
    msg = _("Importate {} righe per {:02d}/{}.").format(len(objects_data), mese, anno)
    if skipped:
        msg += " " + _("{} righe ignorate (colonne insufficienti).").format(skipped)
    messages.success(request, msg)
    return redirect("view-payslip")

@login_required
def delete_payslip_importi(request):
    """
    GET  – mostra il form di selezione mese/anno.
    POST – cancella tutte le righe di payslip_importi per il mese/anno indicati.
    """
    from payroll.models.models import PayslipImporti

    current_year = date.today().year

    def _render(extra=None):
        ctx = {"months": _MONTHS_IT, "current_year": current_year}
        if extra:
            ctx.update(extra)
        return render(request, "payroll/payslip/delete_payslip_importi.html", ctx)

    if request.method == "GET":
        return _render()

    mese_str = request.POST.get("mese", "").strip()
    anno_str = request.POST.get("anno", "").strip()

    try:
        mese = int(mese_str)
        anno = int(anno_str)
        if not (1 <= mese <= 12):
            raise ValueError
    except (ValueError, TypeError):
        messages.error(request, _("Mese o anno non validi."))
        return _render({"mese": mese_str, "anno": anno_str})

    deleted_count, _del_detail = PayslipImporti.objects.filter(mese=mese, anno=anno).delete()
    messages.success(
        request,
        _("Eliminati {} record premi per {:02d}/{}.").format(deleted_count, mese, anno),
    )
    return redirect("view-payslip")

@login_required
def presenze_by_lavoratore(request):
    """
    HTMX partial: tabella dipendenti per un dato mese/anno.
    """
    from django.db.models import Count, Sum as DSum

    try:
        mese = int(request.GET.get("mese", 0))
        anno = int(request.GET.get("anno", 0))
    except (ValueError, TypeError):
        return HttpResponse("")

    search = request.GET.get("search", "").strip()
    qs = PayslipPresenze.objects.filter(mese=mese, anno=anno)
    if search:
        qs = qs.filter(lavoratore__icontains=search)
    rows = (
        qs
        .values("matricola", "lavoratore")
        .annotate(
            n_voci=Count("id"),
            tot_ore=DSum("ore_tot"),
            tot_gg=DSum("gg_tot"),
        )
        .order_by("lavoratore")
    )
    return render(request, "payroll/payslip/presenze_by_lavoratore.html", {
        "rows": rows,
        "mese": mese,
        "anno": anno,
        "search": search,
    })

@login_required
def presenze_lavoratore_rows(request):
    """
    HTMX partial: tutte le colonne di un lavoratore per un dato mese/anno.
    """
    try:
        mese = int(request.GET.get("mese", 0))
        anno = int(request.GET.get("anno", 0))
    except (ValueError, TypeError):
        return HttpResponse("")

    matricola = request.GET.get("matricola", "")
    rows = PayslipPresenze.objects.filter(mese=mese, anno=anno, matricola=matricola).order_by("cod_voce")
    return render(request, "payroll/payslip/presenze_lavoratore_rows.html", {
        "rows": rows,
        "mese": mese,
        "anno": anno,
        "matricola": matricola,
        "lavoratore": rows.first().lavoratore if rows.exists() else matricola,
        "days": list(range(1, 32)),
    })

@login_required
def controllo_cedolini_presenze(request):
    """
    GET  /controllo-cedolini/                  → selettore periodo
    GET  /controllo-cedolini/?mese=2&anno=2026 → pannello configurazione (dizionario + dipendenti)
    POST /controllo-cedolini/                  → esegue il controllo e mostra risultati
    """
    import calendar
    from datetime import date as _date

    periodi = (
        PayslipPresenze.objects
        .values("mese", "anno")
        .distinct()
        .order_by("-anno", "-mese")
    )
    mappings = list(PayslipDizionario.objects.all().order_by("codice_tipo_orario"))

    mese_str = (request.POST.get("mese") or request.GET.get("mese", "")).strip()
    anno_str = (request.POST.get("anno") or request.GET.get("anno", "")).strip()

    regole_controllo = []
    try:
        regole_controllo = list(
            PayslipControlloRegola.objects
            .all()
            .prefetch_related("destinazioni")
            .order_by("direzione", "priorita", "sorgente_valore")
        )
    except Exception:
        # La pagina deve restare accessibile anche prima della migration
        # delle tabelle avanzate di controllo.
        regole_controllo = []

    ctx_base = {
        "periodi": periodi,
        "mappings": mappings,
        "mappings_attivi_count": sum(1 for m in mappings if m.attivo),
        "regole_controllo": regole_controllo,
    }

    if not mese_str or not anno_str:
        return render(request, "payroll/payslip/controllo_cedolini_presenze.html", ctx_base)

    try:
        mese = int(mese_str)
        anno = int(anno_str)
        if not (1 <= mese <= 12):
            raise ValueError
    except (ValueError, TypeError):
        messages.error(request, _("Mese o anno non validi."))
        return render(request, "payroll/payslip/controllo_cedolini_presenze.html", ctx_base)

    num_giorni = calendar.monthrange(anno, mese)[1]
    data_inizio = _date(anno, mese, 1)
    data_fine   = _date(anno, mese, num_giorni)

    # --- Dipendenti da MySQL per il periodo ---
    dipendenti_mysql = []
    try:
        with _pg_conn.cursor() as cur:
            cur.execute(
                """
                SELECT "CODICEPERSONALE", MIN("Descrizione") AS nome
                FROM "_turni_creati"
                WHERE "Data" BETWEEN %s AND %s
                  AND "CODICEPERSONALE" IS NOT NULL AND "CODICEPERSONALE" != ''
                GROUP BY "CODICEPERSONALE"
                ORDER BY nome
                """,
                [data_inizio, data_fine],
            )
            cols = [c[0] for c in cur.description]
            dipendenti_mysql = [dict(zip(cols, row)) for row in cur.fetchall()]
    except Exception as exc:
        messages.error(request, _("Errore lettura turni dal database: {}").format(exc))
        return render(request, "payroll/payslip/controllo_cedolini_presenze.html", {
            **ctx_base, "mese": mese, "anno": anno,
        })

    ctx_config = {
        **ctx_base,
        "mese": mese,
        "anno": anno,
        "dipendenti_mysql": dipendenti_mysql,
    }

    # GET: mostra solo pannello configurazione
    if request.method != "POST":
        return render(request, "payroll/payslip/controllo_cedolini_presenze.html", ctx_config)

    # --- POST: esegui il controllo ---
    all_cod_dip_mysql = [d["CODICEPERSONALE"] for d in dipendenti_mysql]
    selected_dip = request.POST.getlist("dipendenti") or all_cod_dip_mysql

    risultati, mappings_attivi, n_controllati = _build_risultati_for_export(
        mese, anno, selected_dip
    )
    if not mappings_attivi:
        messages.warning(request, _("Nessuna voce attiva nel dizionario."))
        return render(request, "payroll/payslip/controllo_cedolini_presenze.html", {
            **ctx_config,
            "selected_dip": selected_dip,
        })

    return render(request, "payroll/payslip/controllo_cedolini_presenze.html", {
        **ctx_config,
        "selected_dip":    selected_dip,
        "risultati":        risultati,
        "n_controllati":    n_controllati,
        "n_con_diff":       len(risultati),
        "n_discrepanze":    sum(len(r["discrepanze"]) for r in risultati),
        "mappings_attivi":  mappings_attivi,
    })

@login_required
def toggle_dizionario_attivo(request, mapping_id):
    """HTMX: inverte il flag attivo di una riga del dizionario."""
    if request.method != "POST":
        return HttpResponse(status=405)
    m = get_object_or_404(PayslipDizionario, pk=mapping_id)
    m.attivo = not m.attivo
    m.save(update_fields=["attivo"])
    return HttpResponse(status=204)

@login_required
def aggiungi_dizionario(request):
    """Crea una nuova riga nel dizionario di mappatura."""
    if request.method != "POST":
        return HttpResponse(status=405)

    codice_tipo_orario = (request.POST.get("codice_tipo_orario") or "").strip()
    cod_voce = (request.POST.get("cod_voce") or "").strip() or None
    tipo_ora = request.POST.get("tipo_ora", "consuntivo")
    attivo = request.POST.get("attivo") == "1"
    note = (request.POST.get("note") or "").strip() or None

    if not codice_tipo_orario:
        messages.error(request, _("Il campo CODICE_TIPO_ORARIO è obbligatorio."))
    else:
        PayslipDizionario.objects.create(
            codice_tipo_orario=codice_tipo_orario,
            cod_voce=cod_voce,
            tipo_ora=tipo_ora,
            attivo=attivo,
            note=note,
        )
        messages.success(request, _("Voce aggiunta al dizionario."))

    mese = request.POST.get("mese", "")
    anno = request.POST.get("anno", "")
    base = reverse("controllo-cedolini-presenze")
    qs = f"?mese={mese}&anno={anno}" if mese and anno else ""
    return redirect(f"{base}{qs}")

@login_required
def elimina_dizionario(request, mapping_id):
    """Elimina una riga dal dizionario di mappatura."""
    if request.method != "POST":
        return HttpResponse(status=405)
    m = get_object_or_404(PayslipDizionario, pk=mapping_id)
    m.delete()
    messages.success(request, _("Voce eliminata dal dizionario."))
    mese = request.POST.get("mese", "")
    anno = request.POST.get("anno", "")
    base = reverse("controllo-cedolini-presenze")
    qs = f"?mese={mese}&anno={anno}" if mese and anno else ""
    return redirect(f"{base}{qs}")

@login_required
def toggle_regola_controllo_attiva(request, regola_id):
    """HTMX: inverte il flag attiva di una regola avanzata."""
    if request.method != "POST":
        return HttpResponse(status=405)
    reg = get_object_or_404(PayslipControlloRegola, pk=regola_id)
    reg.attivo = not reg.attivo
    reg.save(update_fields=["attivo"])
    return HttpResponse(status=204)

@login_required
def aggiungi_regola_controllo(request):
    """Crea una nuova regola avanzata (tabella payslip_controllo_regole)."""
    if request.method != "POST":
        return HttpResponse(status=405)

    direzione = (request.POST.get("direzione") or "").strip()
    sorgente_valore = (request.POST.get("sorgente_valore") or "").strip()
    modalita = (request.POST.get("modalita") or "").strip()
    no_somma_stesso_giorno = request.POST.get("no_somma_stesso_giorno") == "1"
    attiva = request.POST.get("attiva") == "1"
    note = (request.POST.get("note") or "").strip() or None
    priorita_raw = (request.POST.get("priorita") or "").strip() or "100"
    destinazioni_raw = (request.POST.get("destinazioni") or "").strip()

    if not direzione or not sorgente_valore or not modalita:
        messages.error(request, _("Direzione, sorgente e modalita sono obbligatorie."))
    else:
        try:
            priorita = int(priorita_raw)
        except (ValueError, TypeError):
            priorita = 100

        try:
            regola, created = PayslipControlloRegola.objects.get_or_create(
                direzione=direzione,
                sorgente_valore=sorgente_valore,
                defaults={
                    "modalita": modalita,
                    "no_somma_stesso_giorno": no_somma_stesso_giorno,
                    "attivo": attiva,
                    "priorita": priorita,
                    "note": note,
                },
            )
            if not created:
                regola.modalita = modalita
                regola.no_somma_stesso_giorno = no_somma_stesso_giorno
                regola.attivo = attiva
                regola.priorita = priorita
                regola.note = note
                regola.save(
                    update_fields=[
                        "modalita",
                        "no_somma_stesso_giorno",
                        "attivo",
                        "priorita",
                        "note",
                    ]
                )

            if destinazioni_raw:
                valori = [v.strip() for v in destinazioni_raw.split(",") if v.strip()]
                for val in valori:
                    PayslipControlloRegolaDestinazione.objects.get_or_create(
                        regola=regola,
                        destinazione_valore=val,
                        defaults={"attivo": True},
                    )
            messages.success(request, _("Regola controllo salvata."))
        except Exception as exc:
            messages.error(request, _("Errore nel salvataggio della regola: {}" ).format(exc))

    mese = request.POST.get("mese", "")
    anno = request.POST.get("anno", "")
    base = reverse("controllo-cedolini-presenze")
    qs = f"?mese={mese}&anno={anno}" if mese and anno else ""
    return redirect(f"{base}{qs}")

@login_required
def elimina_regola_controllo(request, regola_id):
    """Elimina una regola avanzata (e le sue destinazioni)."""
    if request.method != "POST":
        return HttpResponse(status=405)
    reg = get_object_or_404(PayslipControlloRegola, pk=regola_id)
    reg.delete()
    messages.success(request, _("Regola eliminata."))
    mese = request.POST.get("mese", "")
    anno = request.POST.get("anno", "")
    base = reverse("controllo-cedolini-presenze")
    qs = f"?mese={mese}&anno={anno}" if mese and anno else ""
    return redirect(f"{base}{qs}")

@login_required
def toggle_regola_destinazione_attiva(request, destinazione_id):
    """HTMX: inverte il flag attiva di una destinazione regola."""
    if request.method != "POST":
        return HttpResponse(status=405)
    d = get_object_or_404(PayslipControlloRegolaDestinazione, pk=destinazione_id)
    d.attivo = not d.attivo
    d.save(update_fields=["attivo"])
    return HttpResponse(status=204)

@login_required
def aggiungi_regola_destinazione(request, regola_id):
    """Aggiunge una destinazione a una regola esistente."""
    if request.method != "POST":
        return HttpResponse(status=405)
    reg = get_object_or_404(PayslipControlloRegola, pk=regola_id)
    val = (request.POST.get("destinazione_valore") or "").strip()
    if not val:
        messages.error(request, _("Destinazione obbligatoria."))
    else:
        PayslipControlloRegolaDestinazione.objects.get_or_create(
            regola=reg,
            destinazione_valore=val,
            defaults={"attivo": True},
        )
        messages.success(request, _("Destinazione aggiunta."))

    mese = request.POST.get("mese", "")
    anno = request.POST.get("anno", "")
    base = reverse("controllo-cedolini-presenze")
    qs = f"?mese={mese}&anno={anno}" if mese and anno else ""
    return redirect(f"{base}{qs}")

@login_required
def elimina_regola_destinazione(request, destinazione_id):
    """Elimina una destinazione da una regola."""
    if request.method != "POST":
        return HttpResponse(status=405)
    d = get_object_or_404(PayslipControlloRegolaDestinazione, pk=destinazione_id)
    d.delete()
    messages.success(request, _("Destinazione eliminata."))
    mese = request.POST.get("mese", "")
    anno = request.POST.get("anno", "")
    base = reverse("controllo-cedolini-presenze")
    qs = f"?mese={mese}&anno={anno}" if mese and anno else ""
    return redirect(f"{base}{qs}")

@login_required
def controllo_cedolini(request):
    """
    Hub page /payroll/controllo-cedolini/ — mostra i tre sotto-menù:
      1. /presenze/  → confronto libro presenze vs turni (già esistente)
      2. /importi/   → controllo voce 429: premi importati vs corpo busta
      3. /acconti/   → da implementare
    """
    # Tronca i breadcrumb in sessione fino a "controllo-cedolini" esclusi,
    # così le sotto-pagine visitate in precedenza non rimangono nel trail.
    crumbs = request.session.get("breadcrumbs", [])
    idx = next(
        (i for i, b in enumerate(crumbs) if b.get("name") == "controllo-cedolini"),
        None,
    )
    if idx is not None:
        request.session["breadcrumbs"] = crumbs[: idx]
        request.session.modified = True

    return render(request, "payroll/payslip/controllo_cedolini.html", {})

@login_required
def controllo_cedolini_importi(request):
    """
    GET  /payroll/controllo-cedolini/importi/ → selettore periodo
    POST /payroll/controllo-cedolini/importi/ → esegue il controllo e mostra risultati

    Incrocia payslip_importi (premi) con payslip_corpo voce 429 per il mese/anno
    selezionato e mostra le discrepanze.
    """
    from payroll.models.models import PayslipCorpo

    periodi_importi = (
        PayslipImporti.objects
        .values("mese", "anno")
        .distinct()
        .order_by("-anno", "-mese")
    )
    periodi_corpo = (
        PayslipCorpo.objects
        .filter(cod_voce=429)
        .values("mese", "anno")
        .distinct()
        .order_by("-anno", "-mese")
    )
    # Unione periodi disponibili
    periodi_set = {(p["mese"], p["anno"]) for p in periodi_importi}
    periodi_set |= {(p["mese"], p["anno"]) for p in periodi_corpo}
    periodi = sorted(periodi_set, key=lambda x: (-x[1], -x[0]))

    mese_str = (request.POST.get("mese") or request.GET.get("mese", "")).strip()
    anno_str = (request.POST.get("anno") or request.GET.get("anno", "")).strip()

    ctx_base = {"periodi": periodi}

    if not mese_str or not anno_str:
        return render(request, "payroll/payslip/controllo_cedolini_importi.html", ctx_base)

    try:
        mese = int(mese_str)
        anno = int(anno_str)
        if not (1 <= mese <= 12):
            raise ValueError
    except (ValueError, TypeError):
        messages.error(request, _("Mese o anno non validi."))
        return render(request, "payroll/payslip/controllo_cedolini_importi.html", ctx_base)

    # --- Dati da payslip_importi ---
    importi_qs = (
        PayslipImporti.objects
        .filter(mese=mese, anno=anno)
        .values("matricola", "badge_id", "neg", "importo")
    )
    importi_by_mat = {}
    for r in importi_qs:
        mat = (r["matricola"] or "").strip()
        if mat:
            importi_by_mat[mat] = {
                "importo": float(r["importo"] or 0) if r["importo"] is not None else 0.0,
                "badge_id": r["badge_id"] or "",
                "neg": r["neg"] or "",
            }

    # --- Dati da payslip_corpo voce 429 ---
    corpo_429_qs = (
        PayslipCorpo.objects
        .filter(mese=mese, anno=anno, cod_voce=429)
        .values("matricola", "cognome", "nome", "importo_ctr_lav")
    )
    corpo_429_by_mat = {}
    for r in corpo_429_qs:
        mat = (r["matricola"] or "").strip()
        if mat:
            nome_completo = f"{r['cognome'] or ''} {r['nome'] or ''}".strip()
            corpo_429_by_mat[mat] = {
                "importo": float(r["importo_ctr_lav"] or 0) if r["importo_ctr_lav"] is not None else 0.0,
                "nome": nome_completo,
            }

    # --- Matricole con almeno una riga in payslip_corpo (qualsiasi voce) ---
    corpo_all_matricole = set(
        (r["matricola"] or "").strip()
        for r in PayslipCorpo.objects
        .filter(mese=mese, anno=anno)
        .values("matricola")
        if (r["matricola"] or "").strip()
    )

    # --- Costruzione risultati ---
    tutte_matricole = set(importi_by_mat.keys()) | set(corpo_429_by_mat.keys())
    risultati = []
    for mat in sorted(tutte_matricole):
        imp = importi_by_mat.get(mat)
        corp = corpo_429_by_mat.get(mat)
        importo_imp = imp["importo"] if imp else None
        importo_corp = corp["importo"] if corp else None
        nome = (corp["nome"] if corp else "") or (imp["badge_id"] if imp else mat)

        if imp and corp:
            # Presente in entrambe le fonti
            delta = round((importo_imp or 0) - (importo_corp or 0), 2)
            stato = "ok"
        elif imp:
            # In payslip_importi ma non in voce 429:
            # distinguiamo se la busta esiste o meno
            delta = None
            if mat in corpo_all_matricole:
                stato = "solo_premi"          # busta presente, ma voce 429 assente
            else:
                stato = "busta_non_presente"  # nessuna riga di busta per questo dipendente
        else:
            # Solo in payslip_corpo voce 429, non in payslip_importi
            delta = None
            stato = "solo_busta"

        risultati.append({
            "matricola":     mat,
            "nome":          nome,
            "neg":           imp["neg"] if imp else "",
            "badge_id":      imp["badge_id"] if imp else "",
            "importo_imp":   importo_imp,
            "importo_corpo": importo_corp,
            "delta":         delta,
            "stato":         stato,
        })

    n_ok   = sum(1 for r in risultati if r["stato"] == "ok")
    n_diff = sum(1 for r in risultati if r["stato"] != "ok")

    return render(request, "payroll/payslip/controllo_cedolini_importi.html", {
        **ctx_base,
        "mese":      mese,
        "anno":      anno,
        "risultati": risultati,
        "n_ok":      n_ok,
        "n_diff":    n_diff,
        "n_totale":  len(risultati),
    })

@login_required
def controllo_cedolini_acconti(request):
    """
    GET  /payroll/controllo-cedolini/acconti/              → selettore periodo
    POST /payroll/controllo-cedolini/acconti/              → esegue il controllo

    Confronta la voce cod_voce=800 (acconto) di payslip_corpo tra il mese
    selezionato e il mese precedente.  Classifica ogni dipendente come:
      - 'nuovo'          : presente nel mese corrente ma assente nel precedente
      - 'sparito'        : presente nel precedente ma assente nel corrente
      - 'cambio_importo' : presente in entrambi ma con importo diverso
      - 'uguale'         : presente in entrambi con stesso importo
    """
    from payroll.models.models import PayslipCorpo

    # Periodi disponibili (solo quelli che hanno almeno una voce 800)
    periodi_qs = (
        PayslipCorpo.objects
        .filter(cod_voce=800)
        .values("mese", "anno")
        .distinct()
        .order_by("-anno", "-mese")
    )
    periodi = [(p["mese"], p["anno"]) for p in periodi_qs]

    mese_str = (request.POST.get("mese") or request.GET.get("mese", "")).strip()
    anno_str = (request.POST.get("anno") or request.GET.get("anno", "")).strip()

    ctx_base = {"periodi": periodi}

    if not mese_str or not anno_str:
        return render(request, "payroll/payslip/controllo_cedolini_acconti.html", ctx_base)

    try:
        mese = int(mese_str)
        anno = int(anno_str)
        if not (1 <= mese <= 12):
            raise ValueError
    except (ValueError, TypeError):
        messages.error(request, _("Mese o anno non validi."))
        return render(request, "payroll/payslip/controllo_cedolini_acconti.html", ctx_base)

    # Mese precedente
    if mese == 1:
        mese_prec, anno_prec = 12, anno - 1
    else:
        mese_prec, anno_prec = mese - 1, anno

    def _fetch_acconti(m, a):
        """Restituisce {matricola: dict} per vo. 800 nel mese/anno m/a."""
        qs = (
            PayslipCorpo.objects
            .filter(mese=m, anno=a, cod_voce=800)
            .values("matricola", "cognome", "nome", "descrizione_voce", "importo_ctr_lav")
            .order_by("cognome", "nome")
        )
        result = {}
        for r in qs:
            mat = (r["matricola"] or "").strip()
            if mat:
                result[mat] = {
                    "matricola":        mat,
                    "cognome":          r["cognome"] or "",
                    "nome":             r["nome"] or "",
                    "descrizione_voce": r["descrizione_voce"] or "",
                    "importo":          float(r["importo_ctr_lav"] or 0),
                }
        return result

    corrente_by_mat  = _fetch_acconti(mese, anno)
    precedente_by_mat = _fetch_acconti(mese_prec, anno_prec)

    # Matricole con almeno una riga in payslip_corpo nel mese corrente (qualsiasi voce)
    corpo_corrente_matricole = set(
        (r["matricola"] or "").strip()
        for r in PayslipCorpo.objects
        .filter(mese=mese, anno=anno)
        .values("matricola")
        if (r["matricola"] or "").strip()
    )

    tutte_matricole = sorted(
        set(corrente_by_mat.keys()) | set(precedente_by_mat.keys()),
        key=lambda m: (
            (corrente_by_mat.get(m) or precedente_by_mat.get(m))["cognome"],
            (corrente_by_mat.get(m) or precedente_by_mat.get(m))["nome"],
        )
    )

    risultati = []
    for mat in tutte_matricole:
        curr = corrente_by_mat.get(mat)
        prev = precedente_by_mat.get(mat)
        ref  = curr or prev   # almeno uno esiste

        importo_corrente  = curr["importo"] if curr else None
        importo_precedente = prev["importo"] if prev else None

        if curr and prev:
            if abs(importo_corrente - importo_precedente) < 0.01:
                stato = "uguale"
            else:
                stato = "cambio_importo"
        elif curr:
            stato = "nuovo"
        else:
            # Presente nel mese precedente ma assente nel corrente:
            # distinguiamo se la busta esiste o meno per il mese corrente
            if mat in corpo_corrente_matricole:
                stato = "sparito"           # busta presente ma voce 800 assente
            else:
                stato = "busta_non_presente"  # nessuna busta paga nel mese corrente

        risultati.append({
            "matricola":          mat,
            "cognome":            ref["cognome"],
            "nome":               ref["nome"],
            "descrizione_voce":   ref["descrizione_voce"],
            "importo_corrente":   importo_corrente,
            "importo_precedente": importo_precedente,
            "delta":              round(importo_corrente - importo_precedente, 2)
                                  if curr and prev else None,
            "stato":              stato,
        })

    n_nuovi           = sum(1 for r in risultati if r["stato"] == "nuovo")
    n_spariti         = sum(1 for r in risultati if r["stato"] == "sparito")
    n_busta_np        = sum(1 for r in risultati if r["stato"] == "busta_non_presente")
    n_cambiati        = sum(1 for r in risultati if r["stato"] == "cambio_importo")
    n_uguali          = sum(1 for r in risultati if r["stato"] == "uguale")

    return render(request, "payroll/payslip/controllo_cedolini_acconti.html", {
        **ctx_base,
        "mese":             mese,
        "anno":             anno,
        "mese_prec":        mese_prec,
        "anno_prec":        anno_prec,
        "risultati":        risultati,
        "n_totale":         len(risultati),
        "n_nuovi":          n_nuovi,
        "n_spariti":        n_spariti,
        "n_busta_np":       n_busta_np,
        "n_cambiati":       n_cambiati,
        "n_uguali":         n_uguali,
    })

@login_required
def controllo_cedolini_malattie(request):
    """
    GET  /payroll/controllo-cedolini/malattie/  → selettore periodo
    POST /payroll/controllo-cedolini/malattie/  → mostra righe presenze per voci MALATTIA

    Legge i cod_voce configurati in payslip_dizionario con codice_tipo_orario='MALATTIA',
    poi estrae da payslip_presenze le righe corrispondenti per il periodo selezionato,
    raggruppandole per dipendente e mostrando i giorni attivi per ogni voce.
    """
    import calendar
    from payroll.models.models import PayslipPresenze

    # --- Voci MALATTIA dal dizionario ---
    voci_malattia_qs = PayslipDizionario.objects.filter(
        codice_tipo_orario="MALATTIA", cod_voce__isnull=False
    ).values("cod_voce", "note").order_by("cod_voce")
    cod_voci_int = []
    for v in voci_malattia_qs:
        try:
            cod_voci_int.append(int(v["cod_voce"]))
        except (TypeError, ValueError):
            pass

    # --- Periodi disponibili ---
    periodi_qs = (
        PayslipPresenze.objects
        .filter(cod_voce__in=cod_voci_int)
        .values("mese", "anno")
        .distinct()
        .order_by("-anno", "-mese")
    ) if cod_voci_int else []
    periodi = [(p["mese"], p["anno"]) for p in periodi_qs]

    mese_str = (request.POST.get("mese") or request.GET.get("mese", "")).strip()
    anno_str = (request.POST.get("anno") or request.GET.get("anno", "")).strip()

    ctx_base = {"periodi": periodi, "n_voci_malattia": len(cod_voci_int)}

    if not mese_str or not anno_str:
        return render(request, "payroll/payslip/controllo_cedolini_malattie.html", ctx_base)

    try:
        mese = int(mese_str)
        anno = int(anno_str)
        if not (1 <= mese <= 12):
            raise ValueError
    except (ValueError, TypeError):
        messages.error(request, _("Mese o anno non validi."))
        return render(request, "payroll/payslip/controllo_cedolini_malattie.html", ctx_base)

    num_giorni = calendar.monthrange(anno, mese)[1]
    day_cols = [f"day_{i}" for i in range(1, 32)]

    qs = (
        PayslipPresenze.objects
        .filter(mese=mese, anno=anno, cod_voce__in=cod_voci_int)
        .values("matricola", "lavoratore", "cod_voce", "desc_voce", "aliq_voce", *day_cols)
        .order_by("lavoratore", "matricola", "cod_voce")
    )

    # Raggruppa per dipendente, converti i giorni in lista compatta dei soli attivi
    from collections import OrderedDict
    import datetime
    _GIORNI_ITA = ["Lu", "Ma", "Me", "Gi", "Ve", "Sa", "Do"]
    dip_dict: dict = OrderedDict()
    for r in qs:
        mat = (r["matricola"] or "").strip()
        lav = (r["lavoratore"] or "").strip()
        key = mat or lav
        if key not in dip_dict:
            dip_dict[key] = {"matricola": mat, "lavoratore": lav, "voci": []}

        # Strip compatta: solo giorni con valore > 0
        giorni_attivi = []
        strip = []  # 31 dict {val, we}, per la barra visiva
        for g in range(1, num_giorni + 1):
            val = r.get(f"day_{g}")
            v = float(val) if val is not None else 0.0
            wd = datetime.date(anno, mese, g).weekday()
            is_we = wd >= 5
            strip.append({"val": round(v, 2) if v > 0 else None, "we": is_we})
            if v > 0:
                giorni_attivi.append({
                    "g": g,
                    "v": round(v, 2),
                    "dow": _GIORNI_ITA[wd],
                    "weekend": is_we,
                })

        cod_voce_val = r["cod_voce"]
        aliq_raw = float(r["aliq_voce"]) if r["aliq_voce"] is not None else None
        # Voce 909: aliquota sempre 100% indipendentemente dal valore del campo
        if cod_voce_val == 909:
            aliq_display = 100.0
        else:
            aliq_display = aliq_raw

        dip_dict[key]["voci"].append({
            "cod_voce":  cod_voce_val,
            "desc_voce": r["desc_voce"] or "",
            "aliq_voce": aliq_display,
            "n_giorni":  len(giorni_attivi),
            "giorni":    giorni_attivi,   # solo i giorni attivi
            "strip":     strip,           # lista 1..num_giorni per barra visiva
        })

    # Ordina le voci di ogni dipendente per aliquota decrescente (None in fondo)
    for dip in dip_dict.values():
        dip["voci"].sort(key=lambda v: v["aliq_voce"] if v["aliq_voce"] is not None else -1, reverse=True)

    risultati = list(dip_dict.values())
    n_giorni_tot = sum(v["n_giorni"] for d in risultati for v in d["voci"])

    return render(request, "payroll/payslip/controllo_cedolini_malattie.html", {
        **ctx_base,
        "mese":         mese,
        "anno":         anno,
        "num_giorni":   num_giorni,
        "risultati":    risultati,
        "n_dipendenti": len(risultati),
        "n_giorni_tot": n_giorni_tot,
    })

def _build_risultati_for_export(mese, anno, selected_dip):
    """
    Ricostruisce i risultati del controllo per un dato periodo e lista dipendenti.
    Restituisce (risultati, mappings_attivi, n_controllati) nel formato usato dalla view principale.

    Gestisce:
    - fallback standard da payslip_dizionario (utile per i casi 1:1)
    - regole avanzate ANY/SUM da payslip_controllo_regole (+ destinazioni)
      per i casi non 1:1.
    """
    import calendar
    from datetime import date as _date
    from collections import defaultdict as _defaultdict

    num_giorni = calendar.monthrange(anno, mese)[1]
    data_inizio = _date(anno, mese, 1)
    data_fine   = _date(anno, mese, num_giorni)

    mappings_attivi = [
        m
        for m in PayslipDizionario.objects.filter(attivo=True).order_by("codice_tipo_orario")
        if m.cod_voce
    ]
    if not mappings_attivi or not selected_dip:
        return [], mappings_attivi, 0

    # Raggruppa per codice_tipo_orario → lista di mapping (es. ROL → [0303, 0336]).
    tipo_orario_to_mappings: dict = _defaultdict(list)
    cod_voce_to_tipi: dict = _defaultdict(set)
    for _m in mappings_attivi:
        tipo_orario_to_mappings[_m.codice_tipo_orario].append(_m)
        try:
            cod_voce_to_tipi[int(str(_m.cod_voce).strip())].add(_m.codice_tipo_orario)
        except (TypeError, ValueError):
            continue

    cod_voce_int_set = set()
    for m in mappings_attivi:
        try:
            cod_voce_int_set.add(int(str(m.cod_voce).strip()))
        except (TypeError, ValueError):
            continue

    # mappa tipo_orario → usa_prev
    tipo_ora_map = {m.codice_tipo_orario: (m.tipo_ora == "previsionale") for m in mappings_attivi}

    # Regole avanzate ANY/SUM (se le tabelle non esistono ancora, fallback automatico).
    regole_attive = []
    try:
        regole_attive = list(
            PayslipControlloRegola.objects
            .filter(attivo=True)
            .prefetch_related("destinazioni")
            .order_by("direzione", "priorita", "sorgente_valore")
        )
    except Exception:
        regole_attive = []

    app_to_ced_rules = {}
    ced_to_app_rules = {}

    for reg in regole_attive:
        dest_vals = [
            (d.destinazione_valore or "").strip()
            for d in reg.destinazioni.all() if d.attivo and (d.destinazione_valore or "").strip()
        ]
        if not dest_vals:
            continue

        src = (reg.sorgente_valore or "").strip()
        if not src:
            continue

        if reg.direzione == PayslipControlloRegola.DIR_APP_TO_CED:
            cods = []
            for v in dest_vals:
                try:
                    cv = int(v)
                    cods.append(cv)
                    cod_voce_int_set.add(cv)
                except ValueError:
                    continue
            if cods:
                app_to_ced_rules[src] = {
                    "modalita": reg.modalita,
                    "no_somma_stesso_giorno": bool(reg.no_somma_stesso_giorno),
                    "dest_cod_voci": cods,
                }

        elif reg.direzione == PayslipControlloRegola.DIR_CED_TO_APP:
            try:
                cod_src = int(src)
            except ValueError:
                continue
            cod_voce_int_set.add(cod_src)
            ced_to_app_rules[cod_src] = {
                "modalita": reg.modalita,
                "no_somma_stesso_giorno": bool(reg.no_somma_stesso_giorno),
                "dest_tipi": dest_vals,
            }

    day_cols = [f"day_{i}" for i in range(1, num_giorni + 1)]

    dipendenti_qs = (
        PayslipPresenze.objects
        .filter(mese=mese, anno=anno, cod_dip__in=selected_dip)
        .exclude(cod_dip__isnull=True).exclude(cod_dip="")
        .values("cod_dip", "lavoratore", "matricola")
        .distinct().order_by("lavoratore")
    )
    dipendenti_list = list(dipendenti_qs)
    all_cod_dip = [d["cod_dip"] for d in dipendenti_list]
    if not all_cod_dip:
        return [], mappings_attivi, 0

    presenze_qs = (
        PayslipPresenze.objects
        .filter(mese=mese, anno=anno,
                cod_voce__in=list(cod_voce_int_set),
                cod_dip__in=all_cod_dip)
        .values("cod_dip", "cod_voce", *day_cols)
    )
    # Se esistono più righe con stesso (cod_dip, cod_voce) sommiamo i valori giornalieri
    presenze_idx = {}
    for r in presenze_qs:
        key = (r["cod_dip"], r["cod_voce"])
        if key not in presenze_idx:
            presenze_idx[key] = dict(r)
        else:
            for dc in day_cols:
                presenze_idx[key][dc] = (
                    float(presenze_idx[key].get(dc) or 0) +
                    float(r.get(dc) or 0)
                )

    # Query PostgreSQL senza filtro tipo → tutti i turni del periodo
    ph_dip = ",".join(["%s"] * len(selected_dip))
    sql = f"""
        SELECT
            "CODICEPERSONALE", "CODICE_TIPO_ORARIO",
            EXTRACT(DAY FROM "Data")::INTEGER AS giorno,
            SUM(CASE WHEN "Ora_Cons_Inizio" IS NOT NULL AND "Ora_Cons_Fine" IS NOT NULL
                          AND "Ora_Cons_Fine" > "Ora_Cons_Inizio"
                     THEN EXTRACT(EPOCH FROM ("Ora_Cons_Fine" - "Ora_Cons_Inizio")) / 3600.0
                     ELSE 0 END) AS ore_cons,
            SUM(CASE WHEN "Ora_Prev_Inizio" IS NOT NULL AND "Ora_Prev_Fine" IS NOT NULL
                          AND "Ora_Prev_Fine" > "Ora_Prev_Inizio"
                     THEN EXTRACT(EPOCH FROM ("Ora_Prev_Fine" - "Ora_Prev_Inizio")) / 3600.0
                     ELSE 0 END) AS ore_prev
        FROM "_turni_creati"
        WHERE "CODICEPERSONALE" IN ({ph_dip})
          AND "Data" BETWEEN %s AND %s
        GROUP BY "CODICEPERSONALE", "CODICE_TIPO_ORARIO", EXTRACT(DAY FROM "Data")
    """
    params = selected_dip + [data_inizio, data_fine]

    turni_idx = {}
    turni_per_day: dict = {}
    with _pg_conn.cursor() as cur:
        cur.execute(sql, params)
        cols = [c[0] for c in cur.description]
        for row in cur.fetchall():
            r = dict(zip(cols, row))
            cp = r["CODICEPERSONALE"]
            tipo = r["CODICE_TIPO_ORARIO"]
            g = int(r["giorno"])
            oc = float(r["ore_cons"] or 0)
            op = float(r["ore_prev"] or 0)
            turni_idx[(cp, tipo, g)] = (oc, op)
            turni_per_day.setdefault((cp, g), []).append((tipo, oc, op))

    def _tipo_effettivo_exp(cod_dip, giorno, usa_prev):
        righe = turni_per_day.get((cod_dip, giorno), [])
        if not righe:
            return "—"
        nomi = [tipo for tipo, oc, op in righe
                if (op if tipo_ora_map.get(tipo, usa_prev) else oc) > 0]
        return ", ".join(nomi) if nomi else righe[0][0]

    def _ore_turno(cod_dip, tipo_orario, giorno):
        ore_cons_m, ore_prev_m = turni_idx.get((cod_dip, tipo_orario, giorno), (0.0, 0.0))
        return ore_prev_m if tipo_ora_map.get(tipo_orario, False) else ore_cons_m

    def _ore_cod_voce(cod_dip, cod_voce_int, giorno):
        return float((presenze_idx.get((cod_dip, cod_voce_int)) or {}).get(f"day_{giorno}") or 0)

    SOGLIA = 0.05
    risultati = []
    n_controllati = 0

    # Fallback legacy: tipi non coperti da regola APP_TO_CED
    fallback_tipi = [
        t for t in tipo_orario_to_mappings.keys()
        if t not in app_to_ced_rules
    ]

    for dip in dipendenti_list:
        cod_dip = dip["cod_dip"]
        discrepanze_dip = []

        # 1) Regole direzionali APP_TO_CED (tipo_orario -> cod_voce)
        for tipo_orario, reg in app_to_ced_rules.items():
            cod_voci = reg["dest_cod_voci"]
            cod_voce_label = " o ".join(f"{cv:04d}" for cv in cod_voci)
            for giorno in range(1, num_giorni + 1):
                ore_app = _ore_turno(cod_dip, tipo_orario, giorno)
                ore_dest = [_ore_cod_voce(cod_dip, cv, giorno) for cv in cod_voci]

                if reg["modalita"] == PayslipControlloRegola.MOD_SUM:
                    ore_ced = sum(ore_dest)
                    valido = abs(ore_app - ore_ced) <= SOGLIA
                else:
                    ore_ced = max(ore_dest) if ore_dest else 0.0
                    valido = any(abs(ore_app - v) <= SOGLIA for v in ore_dest)

                if ore_app <= SOGLIA and ore_ced <= SOGLIA:
                    continue

                if not valido:
                    discrepanze_dip.append({
                        "codice_tipo_orario": f"{tipo_orario} ({reg['modalita']})",
                        "tipo_effettivo_app": _tipo_effettivo_exp(cod_dip, giorno, tipo_ora_map.get(tipo_orario, False)),
                        "cod_voce": cod_voce_label,
                        "desc_voce_cedolino": cod_voce_label,
                        "giorno": giorno,
                        "ore_cedolino": round(ore_ced, 2),
                        "ore_turni": round(ore_app, 2),
                        "delta": round(ore_ced - ore_app, 2),
                    })

        # 2) Fallback legacy per i tipi non coperti da regola.
        for tipo_orario in fallback_tipi:
            tipo_mappings = tipo_orario_to_mappings[tipo_orario]
            cod_voci = []
            for m in tipo_mappings:
                try:
                    cod_voci.append(int(str(m.cod_voce).strip()))
                except (TypeError, ValueError):
                    continue
            if not cod_voci:
                continue

            cod_voce_label = " + ".join(f"{cv:04d}" for cv in cod_voci)
            for giorno in range(1, num_giorni + 1):
                ore_ced = sum(_ore_cod_voce(cod_dip, cv, giorno) for cv in cod_voci)
                ore_app = _ore_turno(cod_dip, tipo_orario, giorno)
                if ore_ced <= SOGLIA and ore_app <= SOGLIA:
                    continue
                if abs(ore_ced - ore_app) > SOGLIA:
                    discrepanze_dip.append({
                        "codice_tipo_orario": tipo_orario,
                        "tipo_effettivo_app": _tipo_effettivo_exp(cod_dip, giorno, tipo_ora_map.get(tipo_orario, False)),
                        "cod_voce": cod_voce_label,
                        "desc_voce_cedolino": cod_voce_label,
                        "giorno": giorno,
                        "ore_cedolino": round(ore_ced, 2),
                        "ore_turni": round(ore_app, 2),
                        "delta": round(ore_ced - ore_app, 2),
                    })

        # 3) Regole direzionali CED_TO_APP (cod_voce -> tipo_orario)
        for cod_voce_src, reg in ced_to_app_rules.items():
            dest_tipi = reg["dest_tipi"]
            if not dest_tipi:
                continue
            tipo_label = " + ".join(dest_tipi)
            for giorno in range(1, num_giorni + 1):
                ore_ced = _ore_cod_voce(cod_dip, cod_voce_src, giorno)
                ore_dest = [_ore_turno(cod_dip, t, giorno) for t in dest_tipi]

                if reg["modalita"] == PayslipControlloRegola.MOD_SUM:
                    ore_app = sum(ore_dest)
                    valido = abs(ore_ced - ore_app) <= SOGLIA
                else:
                    ore_app = max(ore_dest) if ore_dest else 0.0
                    valido = any(abs(ore_ced - v) <= SOGLIA for v in ore_dest)

                if ore_ced <= SOGLIA and ore_app <= SOGLIA:
                    continue

                if not valido:
                    discrepanze_dip.append({
                        "codice_tipo_orario": f"cod_voce {cod_voce_src:04d} ({reg['modalita']})",
                        "tipo_effettivo_app": tipo_label,
                        "cod_voce": f"{cod_voce_src:04d}",
                        "desc_voce_cedolino": f"{cod_voce_src:04d}",
                        "giorno": giorno,
                        "ore_cedolino": round(ore_ced, 2),
                        "ore_turni": round(ore_app, 2),
                        "delta": round(ore_ced - ore_app, 2),
                    })

        n_controllati += 1

        if discrepanze_dip:
            risultati.append({
                "lavoratore": dip.get("lavoratore") or "",
                "cod_dip":    cod_dip,
                "matricola":  dip.get("matricola") or "",
                "discrepanze": sorted(discrepanze_dip, key=lambda x: (x["giorno"], x["codice_tipo_orario"])),
            })

    return risultati, mappings_attivi, n_controllati

@login_required
def export_controllo_excel(request):
    """Scarica un file .xlsx con le discrepanze del controllo cedolini."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    params = request.POST if request.method == "POST" else request.GET
    try:
        mese = int(params.get("mese", 0))
        anno = int(params.get("anno", 0))
        if not (1 <= mese <= 12):
            raise ValueError
    except (ValueError, TypeError):
        return HttpResponse("Parametri non validi", status=400)

    selected_dip = params.getlist("dipendenti") or None
    if not selected_dip:
        selected_dip = list(
            PayslipPresenze.objects
            .filter(mese=mese, anno=anno)
            .exclude(cod_dip__isnull=True).exclude(cod_dip="")
            .values_list("cod_dip", flat=True).distinct()
        )

    risultati, _, _ = _build_risultati_for_export(mese, anno, selected_dip)

    wb = Workbook()
    ws = wb.active
    ws.title = f"Controllo {mese:02d}-{anno}"

    # Intestazione
    headers = ["DIPENDENTE", "GIORNO SETTIMANA", "DATA", "ORE APP", "TURNO APP", "ORE CEDOLINO", "TURNO CEDOLINO"]
    fill_hdr = PatternFill("solid", fgColor="366092")
    font_hdr = Font(bold=True, color="FFFFFF")
    thin = Side(style="thin")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col_idx, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=h)
        cell.font = font_hdr
        cell.fill = fill_hdr
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border
    ws.row_dimensions[1].height = 20

    fill_warn = PatternFill("solid", fgColor="FFF2CC")
    fill_err  = PatternFill("solid", fgColor="FFDCE1")

    from datetime import date as _date
    giorni_settimana_ita = ["lunedì", "martedì", "mercoledì", "giovedì", "venerdì", "sabato", "domenica"]
    row_num = 2
    for ris in risultati:
        for d in ris["discrepanze"]:
            _wd = _date(anno, mese, d["giorno"]).weekday()
            giorno_sett = giorni_settimana_ita[_wd].capitalize()
            data_str = f"{d['giorno']:02d}/{mese:02d}/{anno}"
            values = [
                ris["lavoratore"],
                giorno_sett,
                data_str,
                d["ore_turni"],
                d.get("tipo_effettivo_app", "—"),
                d["ore_cedolino"],
                d["codice_tipo_orario"],
            ]
            row_fill = fill_warn if d["delta"] > 0 else fill_err
            for col_idx, v in enumerate(values, 1):
                cell = ws.cell(row=row_num, column=col_idx, value=v)
                cell.border = border
                cell.fill = row_fill
                if col_idx in (2, 5, 7):
                    cell.alignment = Alignment(horizontal="center")
            row_num += 1

    # Larghezze colonne
    col_widths = [30, 16, 14, 12, 30, 14, 26]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = f'attachment; filename="controllo_cedolini_{mese:02d}_{anno}.xlsx"'
    wb.save(response)
    return response

@login_required
def export_controllo_docx(request):
    """Scarica un file .docx descrittivo con le discrepanze del controllo cedolini."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    params = request.POST if request.method == "POST" else request.GET
    try:
        mese = int(params.get("mese", 0))
        anno = int(params.get("anno", 0))
        if not (1 <= mese <= 12):
            raise ValueError
    except (ValueError, TypeError):
        return HttpResponse("Parametri non validi", status=400)

    selected_dip = params.getlist("dipendenti") or None
    if not selected_dip:
        selected_dip = list(
            PayslipPresenze.objects
            .filter(mese=mese, anno=anno)
            .exclude(cod_dip__isnull=True).exclude(cod_dip="")
            .values_list("cod_dip", flat=True).distinct()
        )

    risultati, _, _ = _build_risultati_for_export(mese, anno, selected_dip)

    doc = Document()

    # Titolo
    title = doc.add_heading(
        f"Controllo Presenze — {mese:02d}/{anno}", level=1
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    mesi_ita = ["", "Gennaio", "Febbraio", "Marzo", "Aprile", "Maggio", "Giugno",
                "Luglio", "Agosto", "Settembre", "Ottobre", "Novembre", "Dicembre"]
    sub = doc.add_paragraph(
        f"Periodo: {mesi_ita[mese]} {anno}  —  "
        f"{sum(len(r['discrepanze']) for r in risultati)} discrepanze su {len(risultati)} dipendenti"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    from datetime import date as _date
    giorni_settimana_ita = ["lunedì", "martedì", "mercoledì", "giovedì", "venerdì", "sabato", "domenica"]

    if not risultati:
        doc.add_paragraph("✅ Nessuna discrepanza rilevata. Tutti i valori coincidono.")
    else:
        for ris in risultati:
            n = len(ris["discrepanze"])
            p = doc.add_heading(
                f"{ris['lavoratore']}  (cod_dip: {ris['cod_dip']})", level=2
            )
            intro = doc.add_paragraph()
            intro.add_run(
                f"Il dipendente {ris['lavoratore']} ha {n} "
                f"{'discrepanza' if n == 1 else 'discrepanze'}:"
            )
            for d in ris["discrepanze"]:
                _wd = _date(anno, mese, d["giorno"]).weekday()
                data_str = f"{giorni_settimana_ita[_wd]} {d['giorno']:02d}/{mese:02d}/{anno}"
                bullet = doc.add_paragraph(style="List Bullet")
                bullet.add_run(f"il {data_str} ").bold = False
                tipo_app_str = d.get('tipo_effettivo_app', '—')
                bullet.add_run(
                    f"sull'app risultano {d['ore_turni']}h di {tipo_app_str}"
                )
                bullet.add_run(" mentre ")
                bullet.add_run(
                    f"sul cedolino sono segnate {d['ore_cedolino']}h di {d['codice_tipo_orario']}"
                )
                bullet.add_run(".")
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    response = HttpResponse(
        buf.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="controllo_cedolini_{mese:02d}_{anno}.docx"'
    return response

@login_required
def export_acconti_excel(request):
    """Scarica un .xlsx con il confronto acconti (voce 800) tra mese corrente e precedente."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from payroll.models.models import PayslipCorpo

    try:
        mese = int(request.GET.get("mese", 0))
        anno = int(request.GET.get("anno", 0))
        if not (1 <= mese <= 12):
            raise ValueError
    except (ValueError, TypeError):
        return HttpResponse("Parametri non validi", status=400)

    if mese == 1:
        mese_prec, anno_prec = 12, anno - 1
    else:
        mese_prec, anno_prec = mese - 1, anno

    def _fetch(m, a):
        qs = PayslipCorpo.objects.filter(mese=m, anno=a, cod_voce=800).values(
            "matricola", "cognome", "nome", "descrizione_voce", "importo_ctr_lav"
        )
        return {
            (r["matricola"] or "").strip(): {
                "cognome": r["cognome"] or "",
                "nome": r["nome"] or "",
                "descrizione_voce": r["descrizione_voce"] or "",
                "importo": float(r["importo_ctr_lav"] or 0),
            }
            for r in qs if (r["matricola"] or "").strip()
        }

    corrente  = _fetch(mese, anno)
    precedente = _fetch(mese_prec, anno_prec)

    # Matricole con almeno una riga in payslip_corpo nel mese corrente (qualsiasi voce)
    corpo_corrente_matricole = set(
        (r["matricola"] or "").strip()
        for r in PayslipCorpo.objects.filter(mese=mese, anno=anno).values("matricola")
        if (r["matricola"] or "").strip()
    )

    tutte = sorted(
        set(corrente.keys()) | set(precedente.keys()),
        key=lambda m: ((corrente.get(m) or precedente.get(m))["cognome"],
                       (corrente.get(m) or precedente.get(m))["nome"])
    )

    wb = Workbook()
    ws = wb.active
    ws.title = f"Acconti {mese:02d}-{anno}"

    thin = Side(style="thin")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    font_hdr = Font(bold=True, color="FFFFFF")
    fill_hdr = PatternFill("solid", fgColor="366092")

    headers = [
        "MATRICOLA", "COGNOME", "NOME", "DESCRIZIONE VOCE",
        f"IMPORTO {mese_prec:02d}/{anno_prec}",
        f"IMPORTO {mese:02d}/{anno}",
        "DIFFERENZA", "STATO",
    ]
    for ci, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=ci, value=h)
        cell.font = font_hdr
        cell.fill = fill_hdr
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border
    ws.row_dimensions[1].height = 20

    fill_map = {
        "nuovo":              PatternFill("solid", fgColor="C6EFCE"),
        "sparito":            PatternFill("solid", fgColor="FFC7CE"),
        "busta_non_presente": PatternFill("solid", fgColor="D9D9D9"),
        "cambio_importo":     PatternFill("solid", fgColor="FFEB9C"),
        "uguale":             None,
    }
    stato_label = {
        "nuovo":              "Nuovo",
        "sparito":            "Sparito",
        "busta_non_presente": "Busta non presente",
        "cambio_importo":     "Importo cambiato",
        "uguale":             "Uguale",
    }

    for row_num, mat in enumerate(tutte, 2):
        curr = corrente.get(mat)
        prev = precedente.get(mat)
        ref  = curr or prev
        imp_c = curr["importo"] if curr else None
        imp_p = prev["importo"] if prev else None
        if curr and prev:
            stato = "uguale" if abs(imp_c - imp_p) < 0.01 else "cambio_importo"
            delta = round(imp_c - imp_p, 2)
        elif curr:
            stato, delta = "nuovo", None
        else:
            delta = None
            stato = "sparito" if mat in corpo_corrente_matricole else "busta_non_presente"

        values = [
            mat, ref["cognome"], ref["nome"], ref["descrizione_voce"],
            imp_p, imp_c, delta, stato_label[stato],
        ]
        fill = fill_map[stato]
        for ci, v in enumerate(values, 1):
            cell = ws.cell(row=row_num, column=ci, value=v)
            cell.border = border
            if fill:
                cell.fill = fill
            if ci in (5, 6, 7):
                cell.alignment = Alignment(horizontal="right")
            if ci == 8:
                cell.alignment = Alignment(horizontal="center")

    col_widths = [16, 22, 18, 30, 18, 18, 14, 20]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = (
        f'attachment; filename="controllo_acconti_{mese:02d}_{anno}.xlsx"'
    )
    wb.save(response)
    return response

@login_required
def export_acconti_docx(request):
    """Scarica un .docx con il confronto acconti (voce 800) tra mese corrente e precedente."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io
    from payroll.models.models import PayslipCorpo

    try:
        mese = int(request.GET.get("mese", 0))
        anno = int(request.GET.get("anno", 0))
        if not (1 <= mese <= 12):
            raise ValueError
    except (ValueError, TypeError):
        return HttpResponse("Parametri non validi", status=400)

    if mese == 1:
        mese_prec, anno_prec = 12, anno - 1
    else:
        mese_prec, anno_prec = mese - 1, anno

    def _fetch(m, a):
        qs = PayslipCorpo.objects.filter(mese=m, anno=a, cod_voce=800).values(
            "matricola", "cognome", "nome", "descrizione_voce", "importo_ctr_lav"
        )
        return {
            (r["matricola"] or "").strip(): {
                "cognome": r["cognome"] or "",
                "nome": r["nome"] or "",
                "descrizione_voce": r["descrizione_voce"] or "",
                "importo": float(r["importo_ctr_lav"] or 0),
            }
            for r in qs if (r["matricola"] or "").strip()
        }

    corrente  = _fetch(mese, anno)
    precedente = _fetch(mese_prec, anno_prec)

    # Matricole con almeno una riga in payslip_corpo nel mese corrente (qualsiasi voce)
    corpo_corrente_matricole = set(
        (r["matricola"] or "").strip()
        for r in PayslipCorpo.objects.filter(mese=mese, anno=anno).values("matricola")
        if (r["matricola"] or "").strip()
    )

    tutte = sorted(
        set(corrente.keys()) | set(precedente.keys()),
        key=lambda m: ((corrente.get(m) or precedente.get(m))["cognome"],
                       (corrente.get(m) or precedente.get(m))["nome"])
    )

    mesi_ita = ["", "Gennaio", "Febbraio", "Marzo", "Aprile", "Maggio", "Giugno",
                "Luglio", "Agosto", "Settembre", "Ottobre", "Novembre", "Dicembre"]

    doc = Document()
    title = doc.add_heading(f"Controllo Acconti — {mese:02d}/{anno}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        f"Confronto {mesi_ita[mese_prec]} {anno_prec} → {mesi_ita[mese]} {anno}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    nuovi    = [(m, corrente[m])   for m in tutte if m in corrente  and m not in precedente]
    spariti  = [(m, precedente[m]) for m in tutte
                if m in precedente and m not in corrente and m in corpo_corrente_matricole]
    busta_np = [(m, precedente[m]) for m in tutte
                if m in precedente and m not in corrente and m not in corpo_corrente_matricole]
    cambiati = [(m, corrente[m], precedente[m]) for m in tutte
                if m in corrente and m in precedente
                and abs(corrente[m]["importo"] - precedente[m]["importo"]) >= 0.01]

    if nuovi:
        doc.add_heading(f"🟢 Nuovi acconti ({len(nuovi)})", level=2)
        for mat, r in nuovi:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{r['cognome']} {r['nome']} (matr. {mat}): ")
            p.add_run(f"{r['importo']:.2f} €").bold = True
        doc.add_paragraph()

    if spariti:
        doc.add_heading(f"🔴 Acconti spariti ({len(spariti)})", level=2)
        for mat, r in spariti:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{r['cognome']} {r['nome']} (matr. {mat}): ")
            p.add_run(f"{r['importo']:.2f} €").bold = True
        doc.add_paragraph()

    if busta_np:
        doc.add_heading(f"⚫ Busta non presente ({len(busta_np)})", level=2)
        for mat, r in busta_np:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{r['cognome']} {r['nome']} (matr. {mat}): ")
            p.add_run(f"{r['importo']:.2f} €").bold = True
            p.add_run(" — nessuna busta paga presente nel mese corrente")
        doc.add_paragraph()

    if cambiati:
        doc.add_heading(f"🟡 Importo cambiato ({len(cambiati)})", level=2)
        for mat, curr, prev in cambiati:
            delta = round(curr["importo"] - prev["importo"], 2)
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{curr['cognome']} {curr['nome']} (matr. {mat}): ")
            p.add_run(f"{prev['importo']:.2f} €").bold = True
            p.add_run(f" → ")
            p.add_run(f"{curr['importo']:.2f} €").bold = True
            p.add_run(f"  (Δ {delta:+.2f} €)")
        doc.add_paragraph()

    if not nuovi and not spariti and not busta_np and not cambiati:
        doc.add_paragraph("✅ Nessuna variazione rilevata sugli acconti.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    response = HttpResponse(
        buf.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = (
        f'attachment; filename="controllo_acconti_{mese:02d}_{anno}.docx"'
    )
    return response

@login_required
def export_malattie_excel(request):
    """Scarica un .xlsx con il controllo malattie per il periodo selezionato."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    import datetime, calendar

    try:
        mese = int(request.GET.get("mese", 0))
        anno = int(request.GET.get("anno", 0))
        if not (1 <= mese <= 12):
            raise ValueError
    except (ValueError, TypeError):
        return HttpResponse("Parametri non validi", status=400)

    num_giorni = calendar.monthrange(anno, mese)[1]
    day_cols = [f"day_{i}" for i in range(1, num_giorni + 1)]
    _GIORNI_ITA = ["Lu", "Ma", "Me", "Gi", "Ve", "Sa", "Do"]

    # Voci MALATTIA dal dizionario
    cod_voci_malattia = list(
        PayslipDizionario.objects.filter(codice_tipo_orario="MALATTIA", attivo=True)
        .exclude(cod_voce__isnull=True)
        .values_list("cod_voce", flat=True)
        .distinct()
    )
    cod_voci_int = [int(v) for v in cod_voci_malattia]

    qs = list(
        PayslipPresenze.objects
        .filter(mese=mese, anno=anno, cod_voce__in=cod_voci_int)
        .values("matricola", "lavoratore", "cod_voce", "desc_voce", "aliq_voce", *day_cols)
        .order_by("lavoratore", "matricola", "cod_voce")
    )

    wb = Workbook()
    ws = wb.active
    ws.title = f"Malattie {mese:02d}-{anno}"

    thin = Side(style="thin")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    font_hdr = Font(bold=True, color="FFFFFF")
    fill_hdr = PatternFill("solid", fgColor="C00000")
    fill_we = PatternFill("solid", fgColor="FFDCE0")   # weekend con ore
    fill_we_hdr = PatternFill("solid", fgColor="7B0000")

    # Intestazioni: colonne fisse + una colonna per ogni giorno del mese
    fixed_headers = ["MATRICOLA", "LAVORATORE", "COD. VOCE", "DESCRIZIONE", "ALIQUOTA %", "N. GIORNI"]
    day_headers = [str(g) for g in range(1, num_giorni + 1)]
    all_headers = fixed_headers + day_headers

    for ci, h in enumerate(all_headers, 1):
        cell = ws.cell(row=1, column=ci, value=h)
        cell.font = font_hdr
        cell.border = border
        cell.alignment = Alignment(horizontal="center", vertical="center")
        idx = ci - len(fixed_headers)  # 1-based day index
        if ci > len(fixed_headers):
            g = idx
            wd = datetime.date(anno, mese, g).weekday()
            cell.fill = fill_we_hdr if wd >= 5 else fill_hdr
        else:
            cell.fill = fill_hdr
    ws.row_dimensions[1].height = 20

    row_num = 2
    for r in qs:
        mat = (r["matricola"] or "").strip()
        lav = (r["lavoratore"] or "").strip()
        cod_voce_val = r["cod_voce"]
        aliq_raw = float(r["aliq_voce"]) if r["aliq_voce"] is not None else None
        aliq_display = 100.0 if cod_voce_val == 909 else aliq_raw
        n_giorni = 0
        for g in range(1, num_giorni + 1):
            v = float(r.get(f"day_{g}") or 0)
            if v > 0:
                n_giorni += 1

        fixed_vals = [mat, lav, cod_voce_val, r["desc_voce"] or "", aliq_display, n_giorni]
        for ci, val in enumerate(fixed_vals, 1):
            cell = ws.cell(row=row_num, column=ci, value=val)
            cell.border = border
            if ci == 5:
                cell.alignment = Alignment(horizontal="right")
            if ci == 6:
                cell.alignment = Alignment(horizontal="center")

        for g in range(1, num_giorni + 1):
            ci = len(fixed_headers) + g
            v = float(r.get(f"day_{g}") or 0)
            val = round(v, 2) if v > 0 else None
            cell = ws.cell(row=row_num, column=ci, value=val)
            cell.border = border
            cell.alignment = Alignment(horizontal="center")
            wd = datetime.date(anno, mese, g).weekday()
            if val is not None and wd >= 5:
                cell.fill = fill_we

        row_num += 1

    # Larghezze colonne
    col_widths_fixed = [16, 28, 11, 32, 12, 10]
    for i, w in enumerate(col_widths_fixed, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    for g in range(1, num_giorni + 1):
        ws.column_dimensions[get_column_letter(len(fixed_headers) + g)].width = 5

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = (
        f'attachment; filename="controllo_malattie_{mese:02d}_{anno}.xlsx"'
    )
    wb.save(response)
    return response

@login_required
def export_malattie_docx(request):
    """Scarica un .docx con il controllo malattie per il periodo selezionato."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io, datetime, calendar

    try:
        mese = int(request.GET.get("mese", 0))
        anno = int(request.GET.get("anno", 0))
        if not (1 <= mese <= 12):
            raise ValueError
    except (ValueError, TypeError):
        return HttpResponse("Parametri non validi", status=400)

    num_giorni = calendar.monthrange(anno, mese)[1]
    day_cols = [f"day_{i}" for i in range(1, num_giorni + 1)]
    _GIORNI_ITA = ["Lu", "Ma", "Me", "Gi", "Ve", "Sa", "Do"]
    mesi_ita = ["", "Gennaio", "Febbraio", "Marzo", "Aprile", "Maggio", "Giugno",
                "Luglio", "Agosto", "Settembre", "Ottobre", "Novembre", "Dicembre"]

    # Voci MALATTIA dal dizionario
    cod_voci_malattia = list(
        PayslipDizionario.objects.filter(codice_tipo_orario="MALATTIA", attivo=True)
        .exclude(cod_voce__isnull=True)
        .values_list("cod_voce", flat=True)
        .distinct()
    )
    cod_voci_int = [int(v) for v in cod_voci_malattia]

    qs = list(
        PayslipPresenze.objects
        .filter(mese=mese, anno=anno, cod_voce__in=cod_voci_int)
        .values("matricola", "lavoratore", "cod_voce", "desc_voce", "aliq_voce", *day_cols)
        .order_by("lavoratore", "matricola", "cod_voce")
    )

    # Raggruppa per dipendente, stessa logica della view
    from collections import OrderedDict
    dip_dict: dict = OrderedDict()
    for r in qs:
        mat = (r["matricola"] or "").strip()
        lav = (r["lavoratore"] or "").strip()
        key = mat or lav
        if key not in dip_dict:
            dip_dict[key] = {"matricola": mat, "lavoratore": lav, "voci": []}

        cod_voce_val = r["cod_voce"]
        aliq_raw = float(r["aliq_voce"]) if r["aliq_voce"] is not None else None
        aliq_display = 100.0 if cod_voce_val == 909 else aliq_raw

        giorni_attivi = []
        for g in range(1, num_giorni + 1):
            v = float(r.get(f"day_{g}") or 0)
            if v > 0:
                wd = datetime.date(anno, mese, g).weekday()
                giorni_attivi.append({
                    "g": g, "v": round(v, 2),
                    "dow": _GIORNI_ITA[wd], "weekend": wd >= 5,
                })

        dip_dict[key]["voci"].append({
            "cod_voce": cod_voce_val,
            "desc_voce": r["desc_voce"] or "",
            "aliq_voce": aliq_display,
            "giorni": giorni_attivi,
        })

    for dip in dip_dict.values():
        dip["voci"].sort(
            key=lambda v: v["aliq_voce"] if v["aliq_voce"] is not None else -1,
            reverse=True,
        )

    doc = Document()
    title = doc.add_heading(f"Controllo Malattie — {mese:02d}/{anno}", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(f"{mesi_ita[mese]} {anno}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    if dip_dict:
        for dip in dip_dict.values():
            # Intestazione dipendente
            hdr = doc.add_heading(level=2)
            run = hdr.add_run(
                f"{dip['lavoratore'] or dip['matricola']}"
                + (f"  (matr. {dip['matricola']})" if dip["matricola"] else "")
            )
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

            for voce in dip["voci"]:
                aliq_str = f"{voce['aliq_voce']:.2f}%" if voce["aliq_voce"] is not None else "—"
                p_voce = doc.add_paragraph()
                r_label = p_voce.add_run(
                    f"Voce {voce['cod_voce']}  {voce['desc_voce']}  [{aliq_str}]  — "
                    f"{len(voce['giorni'])} giorno/i"
                )
                r_label.bold = True
                r_label.font.size = Pt(10)

                if voce["giorni"]:
                    giorni_str = ",  ".join(
                        f"{gd['dow']} {gd['g']}: {gd['v']:.2f}h"
                        + ("  (WE)" if gd["weekend"] else "")
                        for gd in voce["giorni"]
                    )
                    p_det = doc.add_paragraph(style="List Bullet")
                    p_det.add_run(giorni_str).font.size = Pt(9)
                else:
                    p_det = doc.add_paragraph(style="List Bullet")
                    p_det.add_run("nessun giorno attivo").font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            doc.add_paragraph()
    else:
        doc.add_paragraph("Nessuna malattia trovata per il periodo selezionato.")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    response = HttpResponse(
        buf.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = (
        f'attachment; filename="controllo_malattie_{mese:02d}_{anno}.docx"'
    )
    return response



@login_required
@hx_request_required
@permission_required("payroll.add_allowance")
def hx_create_allowance(request):
    """
    This method is used to render htmx allowance form
    """
    form = forms.AllowanceForm()
    return render(request, "payroll/htmx/form.html", {"form": form})


@login_required
# @hx_request_required
@permission_required("payroll.add_payslip")
def send_slip(request):
    """
    Send payslip method
    """

    email_backend = ConfiguredEmailBackend()
    view = request.GET.get("view")
    payslip_ids = request.GET.getlist("id")

    # payslip_ids = request.GET.get("id")
    payslips = Payslip.objects.filter(id__in=payslip_ids)
    if not getattr(
        email_backend, "dynamic_from_email_with_display_name", None
    ) or not len(email_backend.dynamic_from_email_with_display_name):
        messages.error(request, "Email server is not configured")
        if view:
            return HorillaRedirect(request)
        else:
            return redirect(reverse("payslip-list"))

    result_dict = defaultdict(
        lambda: {"employee_id": None, "instances": [], "count": 0}
    )
    for payslip in payslips:
        employee_id = payslip.employee_id
        result_dict[employee_id]["employee_id"] = employee_id
        result_dict[employee_id]["instances"].append(payslip)
        result_dict[employee_id]["count"] += 1

    mail_thread = MailSendThread(request, result_dict=result_dict, ids=payslip_ids)
    mail_thread.start()
    messages.info(request, "Mail processing")
    if view:
        return HorillaRedirect(request)
    else:
        return redirect(reverse("payslip-list"))


@login_required
@permission_required("payroll.add_allowance")
def add_bonus(request):
    employee_id = request.GET.get("employee_id")
    payslip_id = request.GET.get("payslip_id")
    if not employee_id or not payslip_id:
        return HorillaRedirect(request, message=_("Missing required parameters."))
    if payslip_id != "None" and payslip_id:
        instance = Payslip.find(payslip_id)
        if not instance:
            return HorillaRedirect(request, _("Payslip not found"))
        form = forms.PayslipAllowanceForm(
            initial={"employee_id": employee_id, "date": instance.start_date}
        )
    else:
        form = forms.BonusForm(initial={"employee_id": employee_id})

    if request.method == "POST":
        form = forms.BonusForm(request.POST, initial={"employee_id": employee_id})
        contract = Contract.objects.filter(
            employee_id=employee_id, contract_status="active"
        ).first()
        employee = Employee.objects.filter(id=employee_id).first()
        if form.is_valid():
            form.save()
            messages.success(request, _("Bonus Added"))
            if payslip_id != "None" and payslip_id:
                if contract and contract.contract_start_date <= instance.start_date:

                    new_post_data = QueryDict(mutable=True)
                    new_post_data.update(
                        {
                            "employee_id": instance.employee_id,
                            "start_date": instance.start_date,
                            "end_date": instance.end_date,
                        }
                    )
                    instance.delete()
                    create_payslip(request, new_post_data)
                    payslip = Payslip.objects.filter(
                        employee_id=instance.employee_id,
                        start_date=instance.start_date,
                        end_date=instance.end_date,
                    ).first()
                    return HorillaRedirect(
                        request,
                        redirect_to=reverse(
                            "view-payslip", kwargs={"payslip_id": payslip.id}
                        ),
                    )
                else:
                    messages.warning(
                        request,
                        _(
                            "No active contract found for  {} during this payslip period"
                        ).format(employee),
                    )
            return HorillaRedirect(request)

    return render(
        request,
        "payroll/bonus/form.html",
        {"form": form, "employee_id": employee_id, "payslip_id": payslip_id},
    )


@login_required
@permission_required("payroll.add_deduction")
def add_deduction(request):
    employee_id = request.GET.get("employee_id")
    payslip_id = request.GET.get("payslip_id")
    if not employee_id or not payslip_id:
        return HorillaRedirect(request, message=_("Missing required parameters."))
    instance = Payslip.objects.get(id=payslip_id)

    if request.method == "POST":
        form = forms.PayslipDeductionForm(
            request.POST,
            initial={"employee_id": employee_id, "one_time_date": instance.start_date},
        )
        if form.is_valid():
            # Save the form to create the Deduction instance
            deduction_instance = form.save(commit=False)
            deduction_instance.only_show_under_employee = True
            deduction_instance.save()

            # Now that the Deduction instance is saved, add the related employees
            deduction_instance.specific_employees.set([employee_id])
            deduction_instance.include_active_employees = False
            deduction_instance.save()

            # Now create new payslip by deleting existing payslip
            new_post_data = QueryDict(mutable=True)
            new_post_data.update(
                {
                    "employee_id": instance.employee_id,
                    "start_date": instance.start_date,
                    "end_date": instance.end_date,
                }
            )
            instance.delete()
            create_payslip(request, new_post_data)
            payslip = Payslip.objects.filter(
                employee_id=instance.employee_id,
                start_date=instance.start_date,
                end_date=instance.end_date,
            ).first()

            return HorillaRedirect(
                request,
                redirect_to=reverse("view-payslip", kwargs={"payslip_id": payslip.id}),
            )

    else:
        form = forms.PayslipDeductionForm(
            initial={"employee_id": employee_id, "one_time_date": instance.start_date}
        )

    return render(
        request,
        "payroll/deduction/payslip_deduct.html",
        {"form": form, "employee_id": employee_id, "payslip_id": payslip_id},
    )


@login_required
@permission_required("payroll.view_loanaccount")
def view_loans(request):
    """
    This method is used to render template to disply all the loan records
    """
    records = LoanAccount.objects.all()
    loan = records.filter(type="loan")
    adv_salary = records.filter(type="advanced_salary")
    fine = records.filter(type="fine")

    fine_ids = json.dumps(list(fine.values_list("id", flat=True)))
    loan_ids = json.dumps(list(loan.values_list("id", flat=True)))
    adv_salary_ids = json.dumps(list(adv_salary.values_list("id", flat=True)))
    loan = sortby(request, loan, "sortby")
    adv_salary = sortby(request, adv_salary, "sortby")
    fine = sortby(request, fine, "sortby")
    filter_instance = LoanAccountFilter()
    return render(
        request,
        "payroll/loan/view_loan.html",
        {
            "records": paginator_qry(records, request.GET.get("page")),
            "loan": paginator_qry(loan, request.GET.get("lpage")),
            "adv_salary": paginator_qry(adv_salary, request.GET.get("apage")),
            "fine_ids": fine_ids,
            "loan_ids": loan_ids,
            "adv_salary_ids": adv_salary_ids,
            "fine": paginator_qry(fine, request.GET.get("fpage")),
            "f": filter_instance,
        },
    )


@login_required
@hx_request_required
def create_loan(request):
    """
    This method is used to create and update the loan instance
    """
    instance_id = eval_validate(str(request.GET.get("instance_id")))
    instance = LoanAccount.objects.filter(id=instance_id).first()
    form = forms.LoanAccountForm(instance=instance)
    if request.method == "POST":
        form = forms.LoanAccountForm(request.POST, instance=instance)
        if form.is_valid():
            form.save()
            messages.success(request, "Loan created/updated")
            return HorillaRedirect(request)
    return render(
        request, "payroll/loan/form.html", {"form": form, "instance_id": instance_id}
    )


@login_required
@permission_required("payroll.view_loanaccount")
def view_installments(request):
    """
    View install ments
    """
    loan_id = request.GET.get("loan_id")
    if not loan_id:
        return HorillaRedirect(request, message=_("Missing required parameters."))
    loan = LoanAccount.find(loan_id)
    if not loan:
        return HorillaRedirect(request, message=_("Loan not found."))
    installments = loan.deduction_ids.all()

    requests_ids_json = request.GET.get("instances_ids")
    previous_id, next_id = None, None
    if requests_ids_json:
        requests_ids = json.loads(requests_ids_json)
        previous_id, next_id = closest_numbers(requests_ids, int(loan_id))
    return render(
        request,
        "payroll/loan/installments.html",
        {
            "installments": installments,
            "loan": loan,
            "instances_ids": requests_ids_json,
            "previous": previous_id,
            "next": next_id,
        },
    )


@login_required
@permission_required("payroll.delete_loanaccount")
def delete_loan(request):
    """
    Delete loan
    """
    ids = request.GET.getlist("ids")
    loans = LoanAccount.objects.filter(id__in=ids)
    # This 👇 would'nt trigger the delete method in the model
    # loans.delete()
    for loan in loans:
        if (
            not loan.settled
            and not Payslip.objects.filter(
                installment_ids__in=list(
                    loan.deduction_ids.values_list("id", flat=True)
                )
            ).exists()
        ):
            loan.delete()
            messages.success(request, "Loan account deleted")
        else:
            messages.error(request, "Loan account cannot be deleted")
    if request.headers.get("HX-Request"):
        response = HttpResponse("", status=200)
        response["HX-Trigger"] = json.dumps(
            {"reloadPayrollLoanTabs": {"target": "body"}}
        )
        return response
    return redirect(reverse("view-loan"))


@login_required
@permission_required("payroll.view_loanaccount")
def edit_installment_amount(request):
    loan_id = request.GET.get("loan_id")
    ded_id = request.GET.get("ded_id")
    amount_raw = request.POST.get("amount")
    if not loan_id or not ded_id or not amount_raw:
        return HorillaRedirect(request, message=_("Missing required parameters."))
    try:
        value = float(amount_raw) if amount_raw else 0.0
        if not math.isfinite(value):
            value = 0.0
    except (TypeError, ValueError):
        value = 0.0

    loans = LoanAccount.objects.filter(id=loan_id)
    loan = loans.first()
    if not loan:
        return HorillaRedirect(request, message=_("Loan not found."))
    deductions = loan.deduction_ids.all().order_by("one_time_date")
    deduction = deductions.filter(id=ded_id).first()
    deductions_before = deductions.filter(one_time_date__lt=deduction.one_time_date)
    deductions_after = deductions.filter(one_time_date__gt=deduction.one_time_date)
    total_sum = deductions_before.aggregate(Sum("amount"))["amount__sum"] or 0

    balance_instalment = len(deductions_after) if len(deductions_after) != 0 else 1

    new_installment = (loan.loan_amount - total_sum - value) / balance_instalment
    new_installment = round(new_installment, 2)
    if total_sum + value > loan.loan_amount:
        value = loan.loan_amount - total_sum
        new_installment = 0

    if not deduction.installment_payslip():
        deduction.amount = value
        deduction.save()

        for item in deductions.filter(one_time_date__gt=deduction.one_time_date):
            if new_installment > 0:
                item.amount = new_installment
                item.save()
            else:
                item.delete()
                loan.deduction_ids.remove(item)

        # If there are no deductions after the current one and a new installment amount is calculated,
        if len(deductions_after) == 0 and new_installment != 0:
            date = get_next_month_same_date(deduction.one_time_date)
            installment = create_deductions(loan, new_installment, date)
            loan.deduction_ids.add(installment)

        loans.update(installments=len(loan.deduction_ids.all()))
        messages.success(request, "Installment amount updated successfully")
    else:
        messages.error(request, "Cannot change paid installments ")

    return render(
        request,
        # "cbv/loan/loan_detail_view.html",
        "payroll/loan/installments.html",
        {
            "installments": loan.deduction_ids.all(),
            "loan": loan,
        },
    )


@login_required
@hx_request_required
@permission_required("payroll.view_loanaccount")
def search_loan(request):
    """
    Search loan method
    """
    records = LoanAccountFilter(request.GET).qs
    loan = records.filter(type="loan")
    adv_salary = records.filter(type="advanced_salary")
    fine = records.filter(type="fine")

    fine_ids = json.dumps(list(fine.values_list("id", flat=True)))
    loan_ids = json.dumps(list(loan.values_list("id", flat=True)))
    adv_salary_ids = json.dumps(list(adv_salary.values_list("id", flat=True)))
    loan = sortby(request, loan, "sortby")
    adv_salary = sortby(request, adv_salary, "sortby")
    fine = sortby(request, fine, "sortby")

    data_dict = parse_qs(request.GET.urlencode())
    get_key_instances(LoanAccount, data_dict)
    view = request.GET.get("view")
    template = "payroll/loan/records_card.html"
    if view == "list":
        template = "payroll/loan/records_list.html"
    return render(
        request,
        template,
        {
            "records": paginator_qry(records, request.GET.get("page")),
            "loan": paginator_qry(loan, request.GET.get("lpage")),
            "adv_salary": paginator_qry(adv_salary, request.GET.get("apage")),
            "fine": paginator_qry(fine, request.GET.get("fpage")),
            "fine_ids": fine_ids,
            "loan_ids": loan_ids,
            "adv_salary_ids": adv_salary_ids,
            "filter_dict": data_dict,
            "pd": request.GET.urlencode(),
        },
    )


@login_required
@permission_required("payroll.add_loanaccount")
def asset_fine(request):
    """
    Add asset fine method
    """
    if apps.is_installed("asset"):
        Asset = get_horilla_model_class(app_label="asset", model="asset")
    asset_id = request.GET["asset_id"]
    employee_id = request.GET["employee_id"]
    asset = Asset.objects.get(id=asset_id)
    employee = Employee.objects.get(id=employee_id)
    form = forms.AssetFineForm()
    if request.method == "POST":
        form = forms.AssetFineForm(request.POST)
        if form.is_valid():
            instance = form.save(commit=False)
            instance.employee_id = employee
            instance.type = "fine"
            instance.provided_date = date.today()
            instance.asset_id = asset
            instance.save()
            messages.success(request, _("Asset fine added"))
            return HttpResponse(
                "<script>$('#dynamicCreateModal').toggleClass('oh-modal--show'); $('#reloadMessagesButton').click();</script>"
            )  # 880
    return render(
        request,
        "payroll/asset_fine/form.html",
        {"form": form, "asset_id": asset_id, "employee_id": employee_id},
    )


@login_required
def view_reimbursement(request):
    """
    This method is used to render template to view reimbursements
    """
    reimbursement_exists = False
    if Reimbursement.objects.exists():
        reimbursement_exists = True
    if request.GET:
        filter_object = ReimbursementFilter(request.GET)
    else:
        filter_object = ReimbursementFilter({"status": "requested"})
    requests = filter_own_records(
        request, filter_object.qs, "payroll.view_reimbursement"
    )
    reimbursements = requests.filter(type="reimbursement")
    leave_encashments = requests.filter(type="leave_encashment")
    bonus_encashment = requests.filter(type="bonus_encashment")
    data_dict = {"status": ["requested"]}
    view = request.GET.get("view")
    template = "payroll/reimbursement/view_reimbursement.html"

    return render(
        request,
        template,
        {
            "requests": paginator_qry(requests, request.GET.get("page")),
            "reimbursements": paginator_qry(reimbursements, request.GET.get("rpage")),
            "leave_encashments": paginator_qry(
                leave_encashments, request.GET.get("lpage")
            ),
            "bonus_encashments": paginator_qry(
                bonus_encashment, request.GET.get("bpage")
            ),
            "f": filter_object,
            "pd": request.GET.urlencode(),
            "filter_dict": data_dict,
            "view": view,
            "reimbursement_exists": reimbursement_exists,
        },
    )


@login_required
@hx_request_required
def create_reimbursement(request):
    """
    Create or update a reimbursement entry.
    """
    instance = None
    instance_id = request.GET.get("instance_id")

    if instance_id:
        instance = Reimbursement.objects.filter(id=instance_id).first()

    if request.method == "POST":
        form = forms.ReimbursementForm(request.POST, request.FILES, instance=instance)
        if form.is_valid():
            form.save()
            messages.success(request, "Reimbursement saved successfully")
            return HorillaRedirect(request)
    else:
        form = forms.ReimbursementForm(instance=instance)

    return render(request, "payroll/reimbursement/form.html", {"form": form})


@login_required
@hx_request_required
def search_reimbursement(request):
    """
    This method is used to search/filter reimbursement
    """
    requests = ReimbursementFilter(request.GET).qs
    requests = filter_own_records(request, requests, "payroll.view_reimbursement")
    data_dict = parse_qs(request.GET.urlencode())
    reimbursements = requests.filter(type="reimbursement")
    leave_encashments = requests.filter(type="leave_encashment")
    bonus_encashment = requests.filter(type="bonus_encashment")
    reimbursements_ids = json.dumps(list(reimbursements.values_list("id", flat=True)))
    leave_encashments_ids = json.dumps(
        list(leave_encashments.values_list("id", flat=True))
    )
    bonus_encashment_ids = json.dumps(
        list(bonus_encashment.values_list("id", flat=True))
    )
    reimbursements = sortby(request, reimbursements, "sortby")
    leave_encashments = sortby(request, leave_encashments, "sortby")
    bonus_encashment = sortby(request, bonus_encashment, "sortby")
    view = request.GET.get("view")
    template = "payroll/reimbursement/request_cards.html"
    if view == "list":
        template = "payroll/reimbursement/reimbursement_list.html"
    get_key_instances(Reimbursement, data_dict)

    return render(
        request,
        template,
        {
            "requests": paginator_qry(requests, request.GET.get("page")),
            "reimbursements": paginator_qry(reimbursements, request.GET.get("rpage")),
            "leave_encashments": paginator_qry(
                leave_encashments, request.GET.get("lpage")
            ),
            "bonus_encashments": paginator_qry(
                bonus_encashment, request.GET.get("bpage")
            ),
            "filter_dict": data_dict,
            "pd": request.GET.urlencode(),
            "reimbursements_ids": reimbursements_ids,
            "leave_encashments_ids": leave_encashments_ids,
            "bonus_encashment_ids": bonus_encashment_ids,
        },
    )


@login_required
def get_assigned_leaves(request):
    """
    This method is used to return assigned leaves of the employee
    in Json
    """
    emp_id = request.GET.get("employeeId")
    if not emp_id:
        messages.error(request, "Missing required parameters.")
        return JsonResponse(
            {"error": "Missing required parameters: employeeId"}, status=400
        )
    if apps.is_installed("leave"):
        AvailableLeave = get_horilla_model_class(
            app_label="leave", model="availableleave"
        )

    assigned_leaves = (
        AvailableLeave.objects.filter(
            employee_id__id=request.GET["employeeId"],
            total_leave_days__gte=1,
            leave_type_id__is_encashable=True,
        )
        .values(
            "leave_type_id__name",
            "available_days",
            "carryforward_days",
            "leave_type_id__id",
        )
        .distinct()
    )
    return JsonResponse(list(assigned_leaves), safe=False)


@login_required
@permission_required("payroll.change_reimbursement")
def approve_reimbursements(request):
    """
    This method is used to approve or reject the reimbursement request
    """
    ids = request.GET.getlist("ids")
    status = request.GET.get("status")
    if not status:
        return HorillaRedirect(request, message=_("Missing required parameters."))
    if status == "canceled":
        status = "rejected"
    amount = (
        eval_validate(request.GET.get("amount")) if request.GET.get("amount") else 0
    )
    amount = max(0, amount)
    reimbursements = Reimbursement.objects.filter(id__in=ids)
    if status and len(status):
        for reimbursement in reimbursements:
            if reimbursement.type == "leave_encashment":
                reimbursement.amount = amount
            elif reimbursement.type == "bonus_encashment":
                reimbursement.amount = amount

            emp = reimbursement.employee_id
            reimbursement.status = status
            reimbursement.save()
            if reimbursement.status == "requested":
                if not (messages.get_messages(request)._queued_messages):
                    messages.info(request, _("Please check the data you provided."))
            else:
                messages.success(
                    request,
                    _(f"Request {reimbursement.get_status_display()} successfully"),
                )
        if status == "rejected":
            notify.send(
                request.user.employee_get,
                recipient=emp.employee_user_id,
                verb="Your reimbursement request has been rejected.",
                verb_ar="تم رفض طلب استرداد النفقات الخاص بك.",
                verb_de="Ihr Erstattungsantrag wurde abgelehnt.",
                verb_es="Su solicitud de reembolso ha sido rechazada.",
                verb_fr="Votre demande de remboursement a été rejetée.",
                redirect=reverse("view-reimbursement") + f"?id={reimbursement.id}",
                icon="checkmark",
            )
        else:
            notify.send(
                request.user.employee_get,
                recipient=emp.employee_user_id,
                verb="Your reimbursement request has been approved.",
                verb_ar="تمت الموافقة على طلب استرداد نفقاتك.",
                verb_de="Ihr Rückerstattungsantrag wurde genehmigt.",
                verb_es="Se ha aprobado tu solicitud de reembolso.",
                verb_fr="Votre demande de remboursement a été approuvée.",
                redirect=reverse("view-reimbursement") + f"?id={reimbursement.id}",
                icon="checkmark",
            )
    if request.headers.get("HX-Request"):
        response = HttpResponse("", status=200)
        response["HX-Trigger"] = json.dumps(
            {"reloadPayrollReimbursements": {"target": "body"}}
        )
        return response
    return redirect(reverse("view-reimbursement"))


@login_required
@permission_required("payroll.delete_reimbursement")
def delete_reimbursements(request):
    """
    This method is used to delete the reimbursements
    """
    ids = request.GET.getlist("ids")
    reimbursements = Reimbursement.objects.filter(id__in=ids).select_related(
        "employee_id__employee_user_id"
    )
    recipients = []
    seen_user_ids = set()
    for reimbursement in reimbursements:
        recipient = getattr(reimbursement.employee_id, "employee_user_id", None)
        if recipient and recipient.id not in seen_user_ids:
            recipients.append(recipient)
            seen_user_ids.add(recipient.id)
    reimbursements.delete()
    messages.success(request, "Reimbursements deleted")
    if recipients:
        notify.send(
            request.user.employee_get,
            recipient=recipients,
            verb="Your reimbursement request has been deleted.",
            verb_ar="تم حذف طلب استرداد نفقاتك.",
            verb_de="Ihr Rückerstattungsantrag wurde gelöscht.",
            verb_es="Tu solicitud de reembolso ha sido eliminada.",
            verb_fr="Votre demande de remboursement a été supprimée.",
            redirect="/",
            icon="trash",
        )

    if request.headers.get("HX-Request"):
        response = HttpResponse("", status=200)
        response["HX-Trigger"] = json.dumps(
            {"reloadPayrollReimbursements": {"target": "body"}}
        )
        return response
    return redirect("view-reimbursement")


@login_required
@owner_can_enter("payroll.view_reimbursement", Reimbursement, True)
def reimbursement_individual_view(request, instance_id):
    """
    This method is used to render the individual view of reimbursement object
    """
    reimbursement = Reimbursement.find(instance_id)
    if not reimbursement:
        return HorillaRedirect(request, message=_("Reimbursement request not found."))
    requests_ids_json = request.GET.get("instances_ids")
    if requests_ids_json:
        requests_ids = json.loads(requests_ids_json)
        previous_id, next_id = closest_numbers(requests_ids, instance_id)
    context = {
        "reimbursement": reimbursement,
        "instances_ids": requests_ids_json,
        "previous": previous_id,
        "next": next_id,
    }
    return render(
        request,
        "payroll/reimbursement/reimbursenent_individual.html",
        context,
    )


@login_required
@owner_can_enter("payroll.view_reimbursement", Reimbursement, True)
def reimbursement_attachments(request, instance_id):
    """
    This method is used to render all the attachements under the reimbursement object
    """
    reimbursement = Reimbursement.find(instance_id)
    if not reimbursement:
        return HorillaRedirect(request, message=_("Reimbursement request not found."))
    return render(
        request,
        "payroll/reimbursement/attachments.html",
        {"reimbursement": reimbursement},
    )


@login_required
@owner_can_enter("payroll.delete_reimbursement", Reimbursement, True)
def delete_attachments(request, _reimbursement_id):
    """
    This mehtod is used to delete the attachements
    """
    ids = request.GET.getlist("ids")
    ReimbursementMultipleAttachment.objects.filter(id__in=ids).delete()
    messages.success(request, "Attachment deleted")
    return redirect("view-reimbursement")


@login_required
@hx_request_required
@permission_required("payroll.view_payslip")
def get_contribution_report(request):
    """
    This method is used to get the contribution report
    """
    employee_id = request.GET.get("employee_id")
    contribution_deductions = []
    if employee_id:
        pay_heads = Payslip.objects.filter(employee_id__id=employee_id).values_list(
            "pay_head_data", flat=True
        )
        deductions = []
        for head in pay_heads:
            for deduction in head["gross_pay_deductions"]:
                if deduction.get("deduction_id"):
                    deductions.append(deduction)
            for deduction in head["basic_pay_deductions"]:
                if deduction.get("deduction_id"):
                    deductions.append(deduction)
            for deduction in head["pretax_deductions"]:
                if deduction.get("deduction_id"):
                    deductions.append(deduction)
            for deduction in head["post_tax_deductions"]:
                if deduction.get("deduction_id"):
                    deductions.append(deduction)
            for deduction in head["tax_deductions"]:
                if deduction.get("deduction_id"):
                    deductions.append(deduction)
            for deduction in head["net_deductions"]:
                deductions.append(deduction)

        deductions.sort(key=lambda x: x["deduction_id"])
        grouped_deductions = {
            key: list(group)
            for key, group in groupby(deductions, key=lambda x: x["deduction_id"])
        }

        for deduction_id, group in grouped_deductions.items():
            title = group[0]["title"]
            employee_contribution = sum(item.get("amount", 0) for item in group)
            employer_contribution = sum(
                item.get("employer_contribution_amount", 0) for item in group
            )
            total_contribution = employee_contribution + employer_contribution
            if employer_contribution > 0:
                contribution_deductions.append(
                    {
                        "deduction_id": deduction_id,
                        "title": title,
                        "employee_contribution": employee_contribution,
                        "employer_contribution": employer_contribution,
                        "total_contribution": total_contribution,
                    }
                )
    return render(
        request,
        "payroll/dashboard/contribution.html",
        {"contribution_deductions": contribution_deductions},
    )


def all_deductions(pay_head):

    extracted_items = []

    potential_lists = [
        "basic_pay_deductions",
        "gross_pay_deductions",
        "pretax_deductions",
        "post_tax_deductions",
        "tax_deductions",
        "net_deductions",
    ]

    for list_name in potential_lists:
        if list_name in pay_head.keys():
            for item in pay_head[list_name]:
                if "deduction_id" in item:
                    extracted_items.append(item)

    return extracted_items


@login_required
def payslip_detailed_export_data(request):
    """
    This view create the data for exporting payslip data based on selected fields and filters,
    """
    choices_mapping = {
        "draft": _("Draft"),
        "review_ongoing": _("Review Ongoing"),
        "confirmed": _("Confirmed"),
        "paid": _("Paid"),
    }
    selected_columns = []
    payslips_data = []
    totals = {}
    payslips = PayslipFilter(request.GET).qs
    selected_fields = request.GET.getlist("selected_fields")
    form = forms.PayslipExportColumnForm()

    allowances = Allowance.objects.all()
    deductions = Deduction.objects.all()

    if not selected_fields:
        selected_fields = form.fields["selected_fields"].initial

    for field in forms.excel_columns:
        value, key = field

        if value in selected_fields:
            selected_columns.append((value, key))

    selected_columns += [
        (value.title, value.title)
        for value in allowances.filter(
            one_time_date__isnull=True, include_active_employees=True
        )
    ]
    selected_columns += [
        ("other_allowances", "Other Allowances"),
        ("total_allowances", "Total Allowances"),
    ]

    selected_columns += [
        (value.title, value.title)
        for value in deductions.filter(
            one_time_date__isnull=True,
            include_active_employees=True,
            update_compensation__isnull=True,
        )
    ]
    selected_columns += [
        ("federal_tax", "Federal Tax"),
        ("other_deductions", "Other Deductions"),
        ("total_deductions", "Total Deductions"),
    ]

    allowance_totals = {
        column_name.title: 0
        for column_name in allowances.filter(
            one_time_date__isnull=True,
            include_active_employees=True,
        )
    }

    deduction_totals = {
        column_name.title: 0
        for column_name in deductions.filter(
            one_time_date__isnull=True,
            include_active_employees=True,
            update_compensation__isnull=True,
        )
    }

    other_totals = {
        "Other Allowances": 0,
        "Other Deductions": 0,
        "Total Allowances": 0,
        "Total Deductions": 0,
        "Net Pay": 0,
        "Gross Pay": 0,
        "Federal Tax": 0,
    }

    totals.update(allowance_totals)
    totals.update(deduction_totals)
    totals.update(other_totals)
    for payslip in payslips:
        payslip_data = {}
        other_allowances_sum = 0
        other_deductions_sum = 0
        total_allowance = 0
        total_deduction = 0
        total_federal_tax = 0

        federal_tax = payslip.pay_head_data["federal_tax"]
        total_federal_tax += federal_tax

        allos = payslip.pay_head_data["allowances"]
        deducts = all_deductions(payslip.pay_head_data)

        if allos:
            for allowance in allos:
                if not any(
                    str(allowance["title"]) == str(column_name)
                    for item, column_name in selected_columns
                ):
                    other_allowances_sum += (
                        allowance["amount"] if allowance["amount"] is not None else 0
                    )
                total_allowance += allowance["amount"]

        if deducts:
            for deduction in deducts:
                if not any(
                    str(deduction["title"]) == str(column_name)
                    for item, column_name in selected_columns
                ):
                    other_deductions_sum += (
                        deduction["amount"] if deduction["amount"] is not None else 0
                    )
                total_deduction += deduction["amount"]

        for column_value, column_name in selected_columns:
            nested_attributes = column_value.split("__")
            value = payslip
            for attr in nested_attributes:
                value = getattr(value, attr, None)
                if value is None:
                    break
            data = str(value) if value is not None else ""
            if column_name == "Status":
                data = choices_mapping.get(value, "")

            if isinstance(value, date):
                date_format = request.user.employee_get.get_date_format()
                start_date = datetime.strptime(str(value), "%Y-%m-%d").date()

                for format_name, format_string in settings.HORILLA_DATE_FORMATS.items():
                    if format_name == date_format:
                        data = start_date.strftime(format_string)
            else:
                data = str(value) if value is not None else ""

            if allos:
                for allowance in allos:
                    if str(allowance["title"]) == str(column_name):
                        data = (
                            float(allowance["amount"])
                            if allowance["title"] is not None
                            else 0
                        )

            if deducts:
                for deduction in deducts:
                    if str(deduction["title"]) == str(column_name):
                        data = (
                            float(deduction["amount"])
                            if deduction["title"] is not None
                            else 0
                        )

            payslip_data[column_name] = data
            if column_name in totals:
                try:
                    totals[column_name] += float(data)
                except ValueError:
                    pass

        payslip_data["Other Allowances"] = other_allowances_sum
        payslip_data["Other Deductions"] = other_deductions_sum
        payslip_data["Total Allowances"] = total_allowance
        payslip_data["Total Deductions"] = total_deduction
        payslip_data["Federal Tax"] = federal_tax

        totals["Other Allowances"] += other_allowances_sum
        totals["Other Deductions"] += other_deductions_sum
        totals["Total Allowances"] += total_allowance
        totals["Total Deductions"] += total_deduction
        totals["Federal Tax"] += federal_tax

        payslips_data.append(payslip_data)

    totals_row = {}

    for item, column_name in selected_columns:
        if column_name in totals:
            totals_row[column_name] = totals[column_name]
        else:
            totals_row[column_name] = "-"

    totals_row["Other Allowances"] = totals["Other Allowances"]
    totals_row["Other Deductions"] = totals["Other Deductions"]
    totals_row["Total Allowances"] = totals["Total Allowances"]
    totals_row["Total Deductions"] = totals["Total Deductions"]
    totals_row["Employee"] = "Total"

    payslips_data.append(totals_row)

    return {
        "payslips_data": payslips_data,
        "selected_columns": selected_columns,
        "allowances": list(
            allowances.filter(
                one_time_date__isnull=True,
                include_active_employees=True,
            ).values_list("title", flat=True)
        ),
        "deductions": list(
            deductions.filter(
                one_time_date__isnull=True,
                include_active_employees=True,
                update_compensation__isnull=True,
            ).values_list("title", flat=True)
        ),
    }


@login_required
@permission_required("payroll.change_payslip")
def payslip_detailed_export(request):
    """
    Generate an Excel file for download containing detailed payslip data based on
    filters.

    Args:
        request (HttpRequest): The incoming HTTP request object.

    Returns:
        HttpResponse: A response object with the Excel file as an attachment.
    """

    if request.META.get("HTTP_HX_REQUEST"):
        return render(
            request,
            "payroll/payslip/payslip_export_filter.html",
            {
                "export_column": forms.PayslipExportColumnForm(),
                "export_filter": PayslipFilter(request.GET),
                "report": True,
            },
        )

    export_data = payslip_detailed_export_data(request)
    payslips_data = export_data["payslips_data"]
    selected_columns = export_data["selected_columns"]
    allowances = export_data["allowances"]
    deductions = export_data["deductions"]
    today_date = date.today().strftime("%Y-%m-%d")
    file_name = f"Payslip_excel_{today_date}.xlsx"

    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )
    right_border = Border(right=Side(style="thin"))

    wb = Workbook()
    ws = wb.active
    ws.title = "Payslips"

    header_row = [col_name for _, col_name in selected_columns]
    allowances_header = allowances + ["Other Allowances", "Total Allowances"]
    deductions_header = deductions + [
        "Federal Tax",
        "Other Deductions",
        "Total Deductions",
    ]

    basic_cols = len(header_row) - len(allowances_header) - len(deductions_header)
    allowance_cols = len(allowances_header)
    deduction_cols = len(deductions_header)

    merged_sections = [
        (1, basic_cols, "Employee Details", "0000FF"),
        (basic_cols + 1, basic_cols + allowance_cols, "Allowances", "008000"),
        (
            basic_cols + allowance_cols + 1,
            basic_cols + allowance_cols + deduction_cols,
            "Deductions",
            "FF0000",
        ),
    ]

    bold_cols = [
        1,
        basic_cols + allowance_cols,
        basic_cols + allowance_cols + deduction_cols,
    ]

    for start_col, end_col, title, color in merged_sections:
        ws.merge_cells(
            start_row=1, start_column=start_col, end_row=1, end_column=end_col
        )
        cell = ws.cell(row=1, column=start_col, value=title)
        cell.font = Font(color=color, bold=True)
        cell.alignment = Alignment(horizontal="center")
        cell.border = thin_border

        if end_col <= len(header_row):
            ws.cell(row=1, column=end_col).border = thin_border + right_border
    last_row = len(payslips_data) + 2
    ws.row_dimensions[1].height = 25
    ws.row_dimensions[2].height = 20
    ws.row_dimensions[last_row].height = 25

    subheaders = [
        (header_row[:basic_cols], Font(bold=True, color="0000FF")),
        (allowances_header, Font(bold=True, color="008000")),
        (deductions_header, Font(bold=True, color="FF0000")),
    ]

    col_num = 1
    for subheader, font in subheaders:
        for header in subheader:
            cell = ws.cell(row=2, column=col_num, value=str(header))
            cell.font = font
            cell.alignment = Alignment(horizontal="center")
            cell.border = thin_border
            col_num += 1

    for row_num, payslip_data in enumerate(payslips_data, 3):
        for col_num, header in enumerate(header_row, 1):
            cell = ws.cell(
                row=row_num, column=col_num, value=payslip_data.get(header, "")
            )
            if row_num == last_row:
                cell.font = Font(bold=True, color="800080")
                cell.alignment = Alignment(horizontal="right")
            elif col_num in bold_cols:
                cell.font = Font(bold=True)
            cell.border = thin_border

    for col_num, _ in enumerate(header_row, 1):
        max_length = max(
            len(str(cell.value))
            for cell in ws[get_column_letter(col_num)]
            if cell.value is not None
        )
        ws.column_dimensions[get_column_letter(col_num)].width = max_length + 2

    ws.freeze_panes = ws["B3"]

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )
    response["Content-Disposition"] = f"attachment; filename={file_name}.xlsx"
    wb.save(response)

    return response
